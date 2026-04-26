from __future__ import annotations

import re
import textwrap
from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt

from renderers.blocks import (
    _clean,
    render_education,
    render_experience,
    render_header,
    render_project,
    render_skills,
)
from renderers.theme import load_theme
from schemas.cv_model import CVModel


def render_docx(cv_model: CVModel, template: str = "classic", font_override: str = "") -> BytesIO:
    theme = load_theme(template, font_override=font_override)
    doc = Document()
    margin_cm = float(theme.get("margin_cm", 1.5))
    spacing = float(theme.get("spacing", 4))

    for section in doc.sections:
        margin_inch = margin_cm / 2.54
        section.top_margin = Inches(margin_inch)
        section.bottom_margin = Inches(margin_inch)
        section.left_margin = Inches(margin_inch)
        section.right_margin = Inches(margin_inch)

    normal = doc.styles["Normal"]
    normal.font.name = theme.get("font", "Arial")
    body_size = float(theme.get("size", 11))
    section_size = float(theme.get("section_header_size", body_size))
    normal.font.size = Pt(body_size)
    normal.paragraph_format.space_after = Pt(spacing)

    def _para(text: str, bold: bool = False, size: float | None = None,
              indent: float = 0.0, hanging: float = 0.0):
        p = doc.add_paragraph()
        run = p.add_run(_clean(text))
        run.bold = bold
        run.font.size = Pt(size or body_size)
        if indent:
            p.paragraph_format.left_indent = Inches(indent)
        if hanging:
            p.paragraph_format.first_line_indent = Inches(-hanging)
        return p

    def _section_title(title: str):
        _para(title, bold=True, size=section_size)

    # ── HEADER ──
    for i, hl in enumerate(render_header(cv_model)):
        if i == 0:
            _para(hl, bold=True, size=float(theme.get("header_size", 14)))
        else:
            _para(hl)

    # ── SUMMARY ──
    summary = (cv_model.summary or "").strip()
    if summary:
        _section_title("PROFESSIONAL SUMMARY")
        summary = summary.replace("\r\n", "\n").replace("\r", "\n")
        summary = re.sub(r"\n{2,}", "\n", summary)
        for para in summary.split("\n"):
            para = para.strip()
            if para:
                _para(para)

    # ── EXPERIENCE ──
    if cv_model.experiences:
        _section_title("EXPERIENCE")
        for exp in cv_model.experiences:
            lines = render_experience(exp)
            has_title = bool((getattr(exp, "title", "") or "").strip() or
                             (getattr(exp, "company", "") or "").strip())
            for i, txt in enumerate(lines):
                if txt.startswith("- "):
                    _para(txt, indent=0.2, hanging=0.15)
                elif i == 0 and has_title:
                    _para(txt, bold=True)
                else:
                    _para(txt)

    # ── SKILLS ──
    skill_lines = render_skills(cv_model)
    if skill_lines:
        _section_title("SKILLS")
        for txt in skill_lines:
            _para(txt)

    # ── EDUCATION ──
    if cv_model.education:
        _section_title("EDUCATION")
        for edu in cv_model.education:
            edu_lines = render_education(edu)
            for i, txt in enumerate(edu_lines):
                _para(txt, bold=(i == 0))

    # ── PROJECTS ──
    if cv_model.projects:
        _section_title("PROJECTS")
        for proj in cv_model.projects:
            proj_lines = render_project(proj)
            has_name = bool((getattr(proj, "name", "") or "").strip())
            for i, txt in enumerate(proj_lines):
                if txt.startswith("- "):
                    _para(txt, indent=0.2, hanging=0.15)
                elif i == 0 and has_name:
                    _para(txt, bold=True)
                else:
                    _para(txt)

    # ── CERTIFICATIONS ──
    certs = getattr(cv_model, "certifications", None) or []
    if certs:
        _section_title("CERTIFICATIONS")
        for cert in certs:
            name = _clean(getattr(cert, "name", "") or "")
            issuer = _clean(getattr(cert, "issuer", "") or "")
            date = _clean(getattr(cert, "date", "") or "")
            parts = [p for p in [name, issuer, date] if p]
            _para(" \u2013 ".join(parts))

    # ── LANGUAGES ──
    languages = getattr(cv_model, "languages", None) or []
    if languages:
        _section_title("LANGUAGES")
        lang_parts = [_clean(str(la)) for la in languages if str(la or "").strip()]
        _para(", ".join(lang_parts))

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out
