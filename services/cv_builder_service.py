"""
ATS-Optimized CV Builder Service
Generates professional CVs using OpenAI for content enhancement
and python-docx / fpdf2 for document generation.

ATS Rules enforced:
- Single column layout
- Standard fonts (Arial, Calibri, Times New Roman)
- 10.5-12pt font size
- Standard section headers (Summary, Experience, Education, Skills, etc.)
- Bullet points, not paragraphs
- No graphics, icons, tables, text boxes, multi-column
- Measurable results in experience bullets
- Standard date formats (Mon YYYY – Mon YYYY)
- Skills listed as "Skill – Level" (no star ratings)
- No photo, DOB, marital status
"""

import hashlib
import json
import logging
import os
import re
import threading
import time
import unicodedata
import uuid
from datetime import datetime
from io import BytesIO

# ── Instance identity ──────────────────────────────────────────────────────
# Unique per-process ID so multi-instance clusters can trace which worker
# handled a given request.  Stable within a process lifetime.
INSTANCE_ID = os.getenv("INSTANCE_ID", uuid.uuid4().hex[:12])


def _structured_log(
    _logger: logging.Logger,
    level: int,
    event: str,
    **fields: object,
) -> None:
    """Emit a structured JSON log line with standardised fields."""
    payload = {"event": event, **fields}
    _logger.log(level, json.dumps(payload, default=str, ensure_ascii=False))


from renderers.cache import get_cached, make_cache_key, set_cached
from renderers.pdf_renderer import render_pdf as _render_pdf_blocks
from schemas.cv_model import CVModel, Education, Experience
from schemas.cv_schema import CVSchema
from services import rewrite_service
from services.ats_service import analyze_cv
from services.schema_builder import build_schema
from agents.normalize_agent import get_section_order
from services.section_classifier import canonicalize_section_key, PARSER_VERSION
from services.layout_engine import (
    ATS_DEFAULT_TEMPLATE,
    build_layout_schema,
    layout_to_cv_data,
)

logger = logging.getLogger(__name__)

USE_SCHEMA_PIPELINE = os.getenv("USE_SCHEMA_PIPELINE", "1").lower() in ("1", "true", "yes")

MOCK_SERVICES_ON = os.getenv("MOCK_SERVICES", "").lower() in ("1", "true", "yes")

_AI_TIMEOUT_SECONDS = float(os.getenv("AI_TIMEOUT_SECONDS", "20") or "20")
_AI_MAX_RETRIES = int(os.getenv("AI_MAX_RETRIES", "3") or "3")
_AI_CACHE: dict[str, dict] = {}

# ── Global circuit breaker ────────────────────────────────────────────────
# Limits total concurrent parse/render operations *per worker process*.
# This is intentionally process-local (threading.Semaphore) because it
# protects the CPU/memory of this specific worker.  Cross-cluster request
# throttling is handled at the Redis layer (concurrent slots, rate limits).
_GLOBAL_PARSE_LIMIT = int(os.getenv("GLOBAL_PARSE_LIMIT", "10") or "10")
_global_parse_semaphore = threading.Semaphore(_GLOBAL_PARSE_LIMIT)

# ── Safe / degraded mode ──────────────────────────────────────────────────
# When True, skip classifier + AI review and go straight to raw text render.
# Can be set via env var or triggered automatically by rolling error rate.
# NOTE: _error_timestamps / _safe_mode_auto are process-local.  In a
# multi-worker cluster each worker tracks its own error window.  This is
# acceptable: a spike on one worker degrades only that worker, and a global
# outage will trigger all workers independently.
SAFE_MODE = os.getenv("SAFE_MODE", "").lower() in ("1", "true", "yes")

_ERROR_WINDOW_SECONDS = 300  # 5-minute rolling window
_ERROR_THRESHOLD = 10  # auto-degrade after this many errors
_error_timestamps: list[float] = []
_safe_mode_auto = False
_safe_mode_lock = threading.Lock()


def _record_parse_error() -> None:
    """Record a parse error and auto-enable safe mode if threshold is hit."""
    global _safe_mode_auto
    now = time.time()
    with _safe_mode_lock:
        _error_timestamps.append(now)
        # Prune old entries
        cutoff = now - _ERROR_WINDOW_SECONDS
        while _error_timestamps and _error_timestamps[0] < cutoff:
            _error_timestamps.pop(0)
        if len(_error_timestamps) >= _ERROR_THRESHOLD and not _safe_mode_auto:
            _safe_mode_auto = True
            logger.warning(
                "Auto-safe-mode ENABLED: %d parse errors in %ds window",
                len(_error_timestamps),
                _ERROR_WINDOW_SECONDS,
            )


def _is_safe_mode() -> bool:
    """Return True when the system should operate in degraded mode.

    Auto-recovers when the error window is clear (no recent errors).
    """
    global _safe_mode_auto
    if SAFE_MODE:
        return True
    if _safe_mode_auto:
        with _safe_mode_lock:
            now = time.time()
            cutoff = now - _ERROR_WINDOW_SECONDS
            while _error_timestamps and _error_timestamps[0] < cutoff:
                _error_timestamps.pop(0)
            if len(_error_timestamps) < _ERROR_THRESHOLD:
                _safe_mode_auto = False
                logger.info("Auto-safe-mode RECOVERED: error window clear")
                return False
        return True
    return False


# ── Build identity ────────────────────────────────────────────────────────
# Populated at deploy time so logs can tie back to exact code revision.
BUILD_ID = os.getenv("BUILD_ID", "dev").strip()
GIT_SHA = os.getenv("GIT_SHA", "unknown").strip()
PARSER_BUILD = os.getenv("PARSER_BUILD", "local").strip()

# ── Pipeline feature flags ────────────────────────────────────────────────
# Toggle individual pipeline stages without redeploying.
ENABLE_CLASSIFIER = os.getenv("ENABLE_CLASSIFIER", "1").lower() in ("1", "true", "yes")
ENABLE_AI_REVIEW = os.getenv("ENABLE_AI_REVIEW", "1").lower() in ("1", "true", "yes")
ENABLE_SANITIZER = os.getenv("ENABLE_SANITIZER", "1").lower() in ("1", "true", "yes")
ENABLE_FALLBACK = os.getenv("ENABLE_FALLBACK", "1").lower() in ("1", "true", "yes")


def _config_snapshot() -> dict:
    """Return a snapshot of all runtime config for structured logging."""
    return {
        "parser_version": PARSER_VERSION,
        "safe_mode": _is_safe_mode(),
        "build_id": BUILD_ID,
        "git_sha": GIT_SHA,
        "parser_build": PARSER_BUILD,
        "instance_id": INSTANCE_ID,
        "flags": {
            "classifier": ENABLE_CLASSIFIER,
            "ai_review": ENABLE_AI_REVIEW,
            "sanitizer": ENABLE_SANITIZER,
            "fallback": ENABLE_FALLBACK,
        },
    }


# ── Slow-CV guard thresholds ──────────────────────────────────────────────
_SLOW_WARN_SECONDS = float(os.getenv("SLOW_CV_WARN_SECONDS", "2") or "2")
_SLOW_ABORT_SECONDS = float(os.getenv("SLOW_CV_ABORT_SECONDS", "5") or "5")


def _normalize_text(value: str) -> str:
    if not isinstance(value, str):
        value = str(value or "")
    value = unicodedata.normalize("NFC", value)
    value = value.replace("\r\n", "\n").replace("\r", "\n")
    # Remove stray comment artifacts like /* or */
    value = re.sub(r"/\*|\*/", "", value)
    # Strip markdown bold **text** and italic *text*
    value = re.sub(r"\*\*(.+?)\*\*", r"\1", value)
    value = re.sub(r"(?<!\w)\*(.+?)\*(?!\w)", r"\1", value)
    return value


def _fix_case(word: str) -> str:
    known = {
        "html": "HTML",
        "css": "CSS",
        "js": "JavaScript",
        "java": "Java",
        "c": "C",
        "c++": "C++",
        "c#": "C#",
        "python": "Python",
        "git": "Git",
        "github": "GitHub",
        "tcp": "TCP",
        "rtp": "RTP",
        "rtsp": "RTSP",
    }
    clean = _normalize_text(word).strip()
    return known.get(clean.lower(), clean)


def _dedupe_preserve(items):
    seen = set()
    result = []
    for item in items or []:
        text = _normalize_text(item).strip()
        if not text:
            continue
        key = re.sub(r"\s+", " ", text).lower()
        if key in seen:
            continue
        seen.add(key)
        result.append(text)
    return result


def _sanitize_prompt_text(value: str, max_len: int) -> str:
    text = _normalize_text(value)
    text = text.replace("\x00", "").replace("\ufeff", "")
    text = re.sub(r"```+", "`", text)
    text = re.sub(r"<\/?(script|style)[^>]*>", "", text, flags=re.I)
    text = re.sub(r"\s+", " ", text).strip()
    if len(text) > max_len:
        text = text[:max_len]
    return text


def _parse_json_retry(content: str) -> dict:
    raw = (content or "").strip()
    candidates = [raw]
    if raw.startswith("```"):
        no_fence = re.sub(r"^```(?:json)?\s*", "", raw)
        no_fence = re.sub(r"\s*```$", "", no_fence)
        candidates.append(no_fence.strip())
    match = re.search(r"\{[\s\S]*\}", raw)
    if match:
        candidates.append(match.group(0).strip())

    for candidate in candidates:
        try:
            return json.loads(candidate)
        except Exception:
            continue
    raise ValueError("AI response is not valid JSON")


def _validate_enhanced_schema(payload: dict) -> dict:
    if not isinstance(payload, dict):
        raise ValueError("Enhanced payload must be an object")

    normalized = dict(payload)
    normalized.setdefault("summary", "")
    normalized.setdefault("experiences", [])
    normalized.setdefault("skills_categorized", {})
    normalized.setdefault("education", [])
    normalized.setdefault("certifications", [])
    normalized.setdefault("projects", [])

    if not isinstance(normalized["summary"], str):
        normalized["summary"] = str(normalized["summary"] or "")
    for key in ("experiences", "education", "certifications", "projects"):
        if not isinstance(normalized[key], list):
            normalized[key] = []
    if not isinstance(normalized["skills_categorized"], dict):
        normalized["skills_categorized"] = {}

    return normalized


def _truncate_text(value: str, max_len: int) -> str:
    text = _normalize_text(value).strip()
    if len(text) <= max_len:
        return text
    return text[: max_len - 3].rstrip() + "..."


def _estimate_content_size(cv_data: dict) -> int:
    parts = [
        _normalize_text(cv_data.get("summary", "")),
        " ".join(_dedupe_preserve(cv_data.get("languages", []))),
    ]
    for exp in cv_data.get("experiences", []) or []:
        parts.append(_normalize_text(exp.get("title", "")))
        parts.append(_normalize_text(exp.get("company", "")))
        parts.extend(_dedupe_preserve(exp.get("bullets", [])))
    for edu in cv_data.get("education", []) or []:
        parts.append(_normalize_text(edu.get("degree", "")))
        parts.append(_normalize_text(edu.get("school", "")))
    for proj in cv_data.get("projects", []) or []:
        parts.append(_normalize_text(proj.get("name", "")))
        parts.extend(_dedupe_preserve(proj.get("bullets", [])))
    return len(" ".join(parts))


def _normalize_link_value(raw_value: str) -> str:
    value = _normalize_text(raw_value).strip()
    if not value:
        return ""
    value = value.replace("GitHub:", "").replace("github:", "")
    value = value.replace("LinkedIn:", "").replace("linkedin:", "")
    match = re.search(
        r"(?:https?://)?(?:www\.)?(?:linkedin\.com|github\.com|[A-Za-z0-9.-]+\.[A-Za-z]{2,})(?:/\S*)?", value, re.I
    )
    if match:
        value = match.group(0).strip().rstrip(",.;")
    value = re.sub(r"^(?:linkedin|github|portfolio|website)\s*:\s*", "", value, flags=re.I)
    return value


def _normalize_contact_fields(cv_data: dict) -> dict:
    normalized = dict(cv_data or {})

    email_raw = _normalize_text(normalized.get("email", ""))
    email_match = re.search(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", email_raw, re.I)
    normalized["email"] = email_match.group(0) if email_match else ""

    phone_raw = _normalize_text(normalized.get("phone", ""))
    phone_match = re.search(r"(?:\(?\+?\d[\d()\-. ]{7,}\d)", phone_raw)
    normalized["phone"] = phone_match.group(0).strip() if phone_match else ""

    normalized["linkedin"] = _normalize_link_value(normalized.get("linkedin", ""))

    location_raw = _normalize_text(normalized.get("location", ""))
    location_raw = re.sub(r"^(?:phone|email|linkedin|github)\s*:\s*.*$", "", location_raw, flags=re.I)
    # Strip address/location label prefixes (multilingual)
    location_raw = re.sub(
        r"^\s*(?:adres|address|location|adress[ei]?|direcci[oó]n|ubicaci[oó]n"
        r"|standort|lieu|indirizzo|morada|lokasyon|konum)\s*:\s*",
        "",
        location_raw,
        flags=re.I,
    )
    # GUARD: never blank out location — keep original if stripping emptied it
    cleaned_location = location_raw.strip()
    if cleaned_location:
        normalized["location"] = cleaned_location
    # else: keep whatever was already in normalized["location"]

    return normalized


def _build_header_data(cv_data: dict) -> tuple[str, str, str, list[str]]:
    """Return (name, title, location, contact_parts) from cv_data.

    Header order: Name → Title → Location → Contact line.
    Location is always separate — never mixed into the contact line.
    Contact line never contains the name.
    """
    name = (cv_data.get("full_name") or "").strip()
    title_text = (cv_data.get("title") or "").strip()
    location = (cv_data.get("location") or "").strip()

    # GUARD: Only attempt name rescue when model has NO key fields at all.
    # If full_name, email, phone, or location already exist → skip reconstruction.
    _has_fields = bool(name or (cv_data.get("email") or "").strip() or (cv_data.get("phone") or "").strip() or location)
    if not name and not _has_fields:
        from services.cv_autofix_service import _looks_like_person_name

        # Last-resort rescue from summary first line
        summary = (cv_data.get("summary") or "").strip()
        first_line = summary.split("\n")[0].strip() if summary else ""
        if first_line and _looks_like_person_name(first_line):
            name = first_line

    # Build contact parts — never include name, never include location
    name_lower = name.lower().strip()
    name_words = {w.lower() for w in name.split() if len(w) > 1} if name else set()
    contact_parts: list[str] = []
    for field in ("email", "phone", "linkedin"):
        val = (cv_data.get(field) or "").strip()
        if not val:
            continue
        val_lower = val.lower().strip()
        # Never let name leak into contact line
        if name_lower and val_lower == name_lower:
            continue
        if name_lower and len(name_lower) > 3 and name_lower in val_lower:
            continue
        if name_words and val_lower in name_words:
            continue
        contact_parts.append(val)

    # Add social links to contact line
    for sl in cv_data.get("social_links") or []:
        if isinstance(sl, dict):
            url = (sl.get("url") or "").strip()
            platform = (sl.get("platform") or "").strip()
            if url:
                contact_parts.append(f"{platform}: {url}" if platform else url)

    return name, title_text, location, contact_parts


def _compact_for_one_page(cv_data: dict) -> dict:
    compact = dict(cv_data or {})

    strict_mode = os.getenv("CV_ONE_PAGE_STRICT", "0").lower() in ("1", "true", "yes")
    # GUARD: Default behavior: never drop meaningful user blocks.
    # When strict mode is OFF, return data completely untouched.
    if not strict_mode:
        return compact

    # Keep natural CV blocks when content already fits one-page profile.
    if _estimate_content_size(compact) <= 3300 and len(compact.get("experiences", []) or []) <= 4:
        return compact

    # GUARD: Even in strict mode, preserve all languages (never truncate)
    original_languages = list(compact.get("languages", []))

    compact["summary"] = _truncate_text(compact.get("summary", ""), 600)

    experiences = []
    for exp in compact.get("experiences") or []:
        item = dict(exp or {})
        bullets = _dedupe_preserve(item.get("bullets", []))
        item["bullets"] = [_truncate_text(b, 250) for b in bullets]
        item["title"] = _truncate_text(item.get("title", ""), 90)
        item["company"] = _truncate_text(item.get("company", ""), 90)
        item["location"] = _truncate_text(item.get("location", ""), 60)
        experiences.append(item)
    compact["experiences"] = experiences

    compact["education"] = [dict(e or {}) for e in (compact.get("education") or [])]

    skills_categorized = compact.get("skills_categorized") or {}
    compact_skills = {}
    for index, (category, values) in enumerate(skills_categorized.items()):
        if index >= 12:
            break
        deduped_values = _dedupe_preserve(values if isinstance(values, list) else [str(values)])
        compact_skills[_truncate_text(str(category), 30)] = [_truncate_text(v, 50) for v in deduped_values]
    compact["skills_categorized"] = compact_skills

    certifications = []
    for cert in compact.get("certifications") or []:
        item = dict(cert or {})
        item["name"] = _truncate_text(item.get("name", ""), 110)
        item["issuer"] = _truncate_text(item.get("issuer", ""), 60)
        certifications.append(item)
    compact["certifications"] = certifications

    projects = []
    for proj in compact.get("projects") or []:
        item = dict(proj or {})
        item["name"] = _truncate_text(item.get("name", ""), 90)
        item["description"] = _truncate_text(item.get("description", ""), 150)
        bullets = _dedupe_preserve(item.get("bullets", []))
        item["bullets"] = [_truncate_text(b, 150) for b in bullets]
        projects.append(item)
    compact["projects"] = projects

    # GUARD: restore all languages — never lose language data
    compact["languages"] = _dedupe_preserve(original_languages)

    return compact


# ---------------------------------------------------------------------------
# OpenAI helpers
# ---------------------------------------------------------------------------


def _get_openai_client_and_model():
    if MOCK_SERVICES_ON:
        return None, None
    try:
        from services.ai_client_factory import get_ai_client_and_model

        return get_ai_client_and_model()
    except Exception:
        return None, None


def _enhance_cv_with_ai(cv_data: dict, job_description: str, lang: str = "en") -> dict:
    """Use OpenAI to enhance CV content: rewrite bullets with metrics,
    tailor summary to job description, optimize keyword placement."""

    client, model = _get_openai_client_and_model()
    if not client:
        return _mock_enhance(cv_data, job_description, lang)

    lang_instruction = "Turkish" if lang == "tr" else "English"

    safe_job_description = _sanitize_prompt_text(job_description or "", 3000)
    safe_cv_data_json = _sanitize_prompt_text(
        json.dumps(cv_data, ensure_ascii=False, default=str),
        4000,
    )

    cache_key = f"{lang}:{hash((safe_job_description, safe_cv_data_json))}"
    if cache_key in _AI_CACHE:
        return dict(_AI_CACHE[cache_key])

    prompt = f"""You are an expert ATS CV writer. Enhance the following CV data for an ATS-optimized resume.
Language: Write everything in {lang_instruction}.

Job Description:
{safe_job_description}

CV Data (JSON):
{safe_cv_data_json}

Rules:
1. Rewrite the summary to be 2-3 sentences, tailored to the job description, using keywords from it.
2. For each experience entry, rewrite bullets to:
   - Start with strong action verbs
   - Include measurable results (%, $, numbers) where possible
   - Incorporate relevant keywords from the job description
3. Optimize the skills list: put the most relevant skills first, use "Skill – Level" format (Advanced/Intermediate/Beginner).
4. Keep education as-is but ensure proper formatting.
5. Do NOT add fake data. Only enhance what's provided.

Return a JSON object with these exact keys:
{{
  "summary": "enhanced summary text",
  "experiences": [
    {{
      "title": "Job Title",
      "company": "Company Name",
      "location": "City, Country",
      "start_date": "Mon YYYY",
      "end_date": "Mon YYYY or Present",
      "bullets": ["bullet 1", "bullet 2", ...]
    }}
  ],
  "skills_categorized": {{
    "category_name": ["Skill – Level", ...]
  }},
  "education": [
    {{
      "degree": "Degree Name",
      "school": "School Name",
      "location": "City, Country",
      "start_date": "Mon YYYY",
      "end_date": "Mon YYYY",
      "gpa": "3.8/4.0 or null",
      "field": "Field of Study"
    }}
  ],
  "certifications": [
    {{
      "name": "Cert Name",
      "issuer": "Issuer",
      "date": "Mon YYYY"
    }}
  ],
  "projects": [
    {{
      "name": "Project Name",
      "description": "One line description",
      "bullets": ["bullet 1", ...]
    }}
  ]
}}

Return ONLY valid JSON, no markdown fences."""

    try:
        last_exc = None
        for attempt in range(1, _AI_MAX_RETRIES + 1):
            try:
                resp = client.chat.completions.create(
                    model=model,
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.4,
                    max_tokens=3000,
                    timeout=_AI_TIMEOUT_SECONDS,
                )
                content = (resp.choices[0].message.content or "").strip()
                enhanced = _validate_enhanced_schema(_parse_json_retry(content))
                last_exc = None
                break
            except Exception as exc:
                last_exc = exc
                error_text = str(exc).lower()
                retryable = (
                    "timeout" in error_text
                    or "timed out" in error_text
                    or "rate limit" in error_text
                    or "429" in error_text
                )
                if attempt >= _AI_MAX_RETRIES or not retryable:
                    raise
                time.sleep(min(2 ** (attempt - 1), 5))
        if last_exc is not None:
            raise last_exc

        # Merge back any fields AI didn't return
        for key in ("full_name", "email", "phone", "location", "languages", "linkedin"):
            if key in cv_data and key not in enhanced:
                enhanced[key] = cv_data[key]
        _AI_CACHE[cache_key] = dict(enhanced)
        return enhanced
    except Exception as e:
        logger.warning(f"OpenAI CV enhancement failed: {e}")
        return _mock_enhance(cv_data, job_description, lang)


def _mock_enhance(cv_data: dict, job_description: str, lang: str = "en") -> dict:
    """Fallback: return structured data without AI enhancement."""
    result = dict(cv_data)

    # Ensure required structure exists
    if "summary" not in result or not result["summary"]:
        result["summary"] = cv_data.get("summary", "")

    if "experiences" not in result:
        result["experiences"] = []
    for exp in result["experiences"]:
        if "bullets" not in exp:
            exp["bullets"] = []

    if "skills_categorized" not in result:
        raw_skills = cv_data.get("skills", [])
        if isinstance(raw_skills, list) and raw_skills:
            # Categorize skills automatically for ATS compliance
            categories = {
                "Languages": [],
                "Backend & Frameworks": [],
                "Databases": [],
                "DevOps & Cloud": [],
                "Tools & Platforms": [],
            }
            lang_kw = {
                "python",
                "java",
                "javascript",
                "typescript",
                "c++",
                "c#",
                "go",
                "rust",
                "ruby",
                "php",
                "swift",
                "kotlin",
                "scala",
                "r",
                "matlab",
                "sql",
                "html",
                "css",
                "dart",
                "perl",
                "bash",
                "shell",
                "lua",
                "elixir",
                "haskell",
                "objective-c",
            }
            backend_kw = {
                "django",
                "flask",
                "fastapi",
                "spring",
                "express",
                "nestjs",
                "rails",
                "laravel",
                "react",
                "vue",
                "angular",
                "next",
                "nuxt",
                "svelte",
                "node",
                "asp.net",
                ".net",
                "graphql",
                "rest",
                "grpc",
                "celery",
                "gin",
                "fiber",
                "actix",
            }
            db_kw = {
                "postgresql",
                "postgres",
                "mysql",
                "mongodb",
                "redis",
                "sqlite",
                "oracle",
                "dynamodb",
                "cassandra",
                "elasticsearch",
                "neo4j",
                "mariadb",
                "supabase",
                "firebase",
                "firestore",
                "couchdb",
                "influxdb",
                "mssql",
                "pgvector",
            }
            devops_kw = {
                "docker",
                "kubernetes",
                "k8s",
                "aws",
                "gcp",
                "azure",
                "terraform",
                "ansible",
                "jenkins",
                "ci/cd",
                "github actions",
                "gitlab",
                "linux",
                "nginx",
                "prometheus",
                "grafana",
                "helm",
                "argocd",
                "cloudflare",
                "vercel",
                "heroku",
                "digitalocean",
                "lambda",
                "ec2",
                "s3",
                "ecs",
                "fargate",
            }

            for skill in raw_skills:
                s_lower = skill.lower().strip()
                if s_lower in lang_kw:
                    categories["Languages"].append(skill)
                elif s_lower in backend_kw or any(k in s_lower for k in backend_kw):
                    categories["Backend & Frameworks"].append(skill)
                elif s_lower in db_kw or any(k in s_lower for k in db_kw):
                    categories["Databases"].append(skill)
                elif s_lower in devops_kw or any(k in s_lower for k in devops_kw):
                    categories["DevOps & Cloud"].append(skill)
                else:
                    categories["Tools & Platforms"].append(skill)

            # Only include non-empty categories
            result["skills_categorized"] = {k: v for k, v in categories.items() if v}
            if not result["skills_categorized"]:
                result["skills_categorized"] = {"Technical Skills": raw_skills}
        elif isinstance(raw_skills, dict):
            result["skills_categorized"] = raw_skills
        else:
            result["skills_categorized"] = {}

    if "education" not in result:
        result["education"] = []
    if "certifications" not in result:
        result["certifications"] = []
    if "projects" not in result:
        result["projects"] = []

    return result


# ---------------------------------------------------------------------------
# Language-aware section titles
# ---------------------------------------------------------------------------

_SECTION_TITLES = {
    "en": {
        "summary": "Professional Summary",
        "experience": "Experience",
        "education": "Education",
        "skills": "Skills",
        "certifications": "Certifications",
        "projects": "Projects",
        "languages": "Languages",
        "interests": "Interests",
        "misc": "Other",
    },
    "tr": {
        "summary": "Profesyonel \u00d6zet",
        "experience": "Deneyim",
        "education": "E\u011fitim",
        "skills": "Beceriler",
        "certifications": "Sertifikalar",
        "projects": "Projeler",
        "languages": "Diller",
        "interests": "\u0130lgi Alanlar\u0131",
        "misc": "Di\u011fer",
    },
    "de": {
        "summary": "Zusammenfassung",
        "experience": "Berufserfahrung",
        "education": "Ausbildung",
        "skills": "F\u00e4higkeiten",
        "certifications": "Zertifikate",
        "projects": "Projekte",
        "languages": "Sprachen",
        "interests": "Interessen",
        "misc": "Sonstiges",
    },
    "fr": {
        "summary": "R\u00e9sum\u00e9 Professionnel",
        "experience": "Exp\u00e9rience",
        "education": "Formation",
        "skills": "Comp\u00e9tences",
        "certifications": "Certifications",
        "projects": "Projets",
        "languages": "Langues",
        "interests": "Centres d'Int\u00e9r\u00eat",
        "misc": "Autres",
    },
    "es": {
        "summary": "Resumen Profesional",
        "experience": "Experiencia",
        "education": "Educaci\u00f3n",
        "skills": "Habilidades",
        "certifications": "Certificaciones",
        "projects": "Proyectos",
        "languages": "Idiomas",
        "interests": "Intereses",
        "misc": "Otros",
    },
}


def _section_title(key: str, lang: str = "en") -> str:
    """Return the localised section title for *key* in *lang*."""
    titles = _SECTION_TITLES.get(lang, _SECTION_TITLES["en"])
    return titles.get(key, _SECTION_TITLES["en"].get(key, key.title()))


# ---------------------------------------------------------------------------
# Safety remap — ensure cv_data keys are canonical before rendering
# ---------------------------------------------------------------------------


def _remap_cv_data(cv_data: dict) -> dict:
    """Remap any non-canonical section keys in cv_data so renderers find them."""
    _REMAP = {
        "summary": "summary",
        "experience": "experiences",
        "education": "education",
        "skills": "skills",
        "projects": "projects",
        "certifications": "certifications",
        "languages": "languages",
        "interests": "interests",
        "misc": "misc",
    }
    _SKIP = {
        "full_name",
        "title",
        "email",
        "phone",
        "location",
        "linkedin",
        "summary",
        "experiences",
        "education",
        "skills",
        "skills_categorized",
        "projects",
        "certifications",
        "languages",
        "interests",
        "misc",
        "language",
        "section_titles",
        "format_hints",
        "contact",
        "template",
        "social_links",
    }
    out = dict(cv_data)
    for key in list(out.keys()):
        if key.startswith("_") or key in _SKIP:
            continue
        canonical = canonicalize_section_key(key)
        target = _REMAP.get(canonical, "misc")
        value = out.pop(key)
        if not value:
            continue
        if target == "summary":
            existing = out.get("summary", "")
            extra = value if isinstance(value, str) else " ".join(str(v) for v in value)
            out["summary"] = f"{existing} {extra}".strip() if existing else extra
        else:
            existing = out.get(target) or []
            if isinstance(existing, list) and isinstance(value, list):
                existing.extend(value)
            out[target] = existing
    return out


# ---------------------------------------------------------------------------
# DOCX Generation (ATS-optimized)
# ---------------------------------------------------------------------------


def generate_docx(cv_data: dict, template: str = "classic", lang: str = "en", font_family: str = "") -> BytesIO:
    """Generate an ATS-friendly DOCX file. Returns BytesIO buffer."""
    cv_data = _remap_cv_data(cv_data)
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # -- Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # -- Style helpers
    style = doc.styles["Normal"]
    font = style.font
    font.size = Pt(10)

    # Template-specific fonts and styling
    template_config = {
        "modern": {"font": "Calibri", "accent_color": (0x2B, 0x6C, 0xB0)},
        "executive": {"font": "Times New Roman", "accent_color": (0x8B, 0x45, 0x13)},
        "professional": {"font": "Georgia", "accent_color": (0x2F, 0x4F, 0x4F)},
        "creative": {"font": "Garamond", "accent_color": (0x9B, 0x59, 0xB6)},
        "corporate": {"font": "Cambria", "accent_color": (0x1F, 0x4E, 0x79)},
        "tech": {"font": "Consolas", "accent_color": (0x00, 0x7A, 0xCC)},
        "consulting": {"font": "Book Antiqua", "accent_color": (0x5D, 0x4E, 0x75)},
        "classic": {"font": "Arial", "accent_color": (0x00, 0x00, 0x00)},
    }

    config = template_config.get(template, template_config["classic"])
    # User font override takes precedence over template default
    from renderers.theme import ALLOWED_FONTS

    if font_family and font_family in ALLOWED_FONTS:
        font.name = font_family
    else:
        font.name = config["font"]

    # Keep paragraph rhythm consistent to avoid visual drift between sections.
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    cv_data = _normalize_contact_fields(cv_data)

    name, title_text, location_text, contact_parts = _build_header_data(cv_data)

    if name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name.upper())
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = font.name

    if title_text:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title_text)
        run.font.size = Pt(11)
        run.font.name = font.name

    if location_text:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(location_text)
        run.font.size = Pt(10)
        run.font.name = font.name

    if contact_parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("  |  ".join(contact_parts))
        run.font.size = Pt(10)
        run.font.name = font.name

    def add_section_header(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = font.name

        # Apply template-specific accent color
        color = config["accent_color"]
        if color != (0x00, 0x00, 0x00):  # Not black (classic default)
            run.font.color.rgb = RGBColor(*color)

        # Template-specific styling
        if template in ["corporate", "consulting", "executive"]:
            # Professional templates get thicker border
            border_size = "6"
            border_color = f"{color[0]:02x}{color[1]:02x}{color[2]:02x}"
        elif template in ["creative", "tech"]:
            # Modern templates get colored border
            border_size = "4"
            border_color = f"{color[0]:02x}{color[1]:02x}{color[2]:02x}"
        else:
            # Classic styling
            border_size = "4"
            border_color = "999999"

        # Add separator line
        from docx.oxml.ns import qn

        pPr = p._p.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pBdr.makeelement(
            qn("w:bottom"),
            {
                qn("w:val"): "single",
                qn("w:sz"): border_size,
                qn("w:space"): "1",
                qn("w:color"): border_color,
            },
        )
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_bullet(text_content):
        # Clean text content to avoid special characters
        clean_text = _normalize_text(text_content).strip() if text_content else ""
        if not clean_text:
            return
        # Remove any existing bullet markers
        clean_text = re.sub(r"^[•\-\*]\s*", "", clean_text)

        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(f"• {clean_text}")
        run.font.size = Pt(10.5)
        run.font.name = font.name

    # -- Professional Summary
    summary = _normalize_text(cv_data.get("summary", ""))

    def _docx_summary():
        if not summary:
            return
        add_section_header(_section_title("summary", lang))
        p = doc.add_paragraph()
        run = p.add_run(summary)
        run.font.size = Pt(10.5)
        run.font.name = font.name

    # -- Experience
    experiences = cv_data.get("experiences", [])

    def _docx_experience():
        if not experiences:
            return
        add_section_header(_section_title("experience", lang))
        for exp in experiences:
            # Title + Company line
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(1)

            title_run = p.add_run(_normalize_text(exp.get("title", "")))
            title_run.bold = True
            title_run.font.size = Pt(11)
            title_run.font.name = font.name

            company = _normalize_text(exp.get("company", ""))
            if company:
                sep_run = p.add_run("  —  ")
                sep_run.font.name = font.name
                sep_run.font.size = Pt(11)
                c_run = p.add_run(company)
                c_run.font.size = Pt(11)
                c_run.font.name = font.name

            # Date + Location on a SEPARATE line to prevent overflow
            start = _normalize_text(exp.get("start_date", ""))
            end = _normalize_text(exp.get("end_date", ""))
            location = _normalize_text(exp.get("location", ""))
            meta_parts = []
            if start or end:
                date_str = f"{start} – {end}" if start and end else (start or end)
                meta_parts.append(date_str)
            if location:
                meta_parts.append(location)
            if meta_parts:
                mp = doc.add_paragraph()
                mp.paragraph_format.space_before = Pt(0)
                mp.paragraph_format.space_after = Pt(1)
                mp_run = mp.add_run("  |  ".join(meta_parts))
                mp_run.font.size = Pt(10)
                mp_run.font.name = font.name

            for bullet in _dedupe_preserve(exp.get("bullets", [])):
                if bullet and bullet.strip():
                    add_bullet(bullet)

    # -- Education
    education = cv_data.get("education", [])

    def _docx_education():
        if not education:
            return
        add_section_header(_section_title("education", lang))
        last_gpa = None
        for edu in education:
            # Line 1: Degree (Date) — bold
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(0)

            degree = _normalize_text(edu.get("degree", ""))
            field = _normalize_text(edu.get("field", ""))
            degree_text = f"{degree}" + (f" in {field}" if field else "")

            start = _normalize_text(edu.get("start_date", ""))
            end = _normalize_text(edu.get("end_date", ""))
            if start or end:
                date_str = f"{start} – {end}" if start and end else (start or end)
                degree_text = f"{degree_text} ({date_str})" if degree_text.strip() else date_str

            d_run = p.add_run(degree_text)
            d_run.bold = True
            d_run.font.size = Pt(11)
            d_run.font.name = font.name

            # Line 2: School — normal
            school = _normalize_text(edu.get("school", ""))
            if school:
                sp = doc.add_paragraph()
                sp.paragraph_format.space_before = Pt(0)
                sp.paragraph_format.space_after = Pt(0)
                s_run = sp.add_run(school)
                s_run.font.size = Pt(10.5)
                s_run.font.name = font.name

            # Line 3: Location + GPA
            loc = _normalize_text(edu.get("location", ""))
            gpa = _normalize_text(edu.get("gpa", ""))
            meta_parts = []
            if loc:
                meta_parts.append(loc)
            if gpa:
                normalized_gpa = gpa.replace("GPA:", "").replace("gpa:", "").strip()
                if normalized_gpa and normalized_gpa != last_gpa:
                    meta_parts.append(f"GPA: {normalized_gpa}")
                    last_gpa = normalized_gpa
            if meta_parts:
                mp = doc.add_paragraph()
                mp.paragraph_format.space_before = Pt(0)
                mp.paragraph_format.space_after = Pt(1)
                mp_run = mp.add_run("  |  ".join(meta_parts))
                mp_run.font.size = Pt(10)
                mp_run.font.name = font.name

    # -- Skills
    skills_cat = cv_data.get("skills_categorized", {})

    def _docx_skills():
        if not skills_cat:
            return
        add_section_header(_section_title("skills", lang))
        for category, items in skills_cat.items():
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(1)
            cat_run = p.add_run(f"{_normalize_text(category)}: ")
            cat_run.bold = True
            cat_run.font.size = Pt(10.5)
            cat_run.font.name = font.name
            skills_values = _dedupe_preserve(items if isinstance(items, list) else [str(items)])
            skills_values = [_fix_case(v) for v in skills_values]
            skills_text = ", ".join(skills_values)
            sk_run = p.add_run(skills_text)
            sk_run.font.size = Pt(10.5)
            sk_run.font.name = font.name

    # -- Certifications
    certs = cv_data.get("certifications", [])

    def _docx_certifications():
        if not certs:
            return
        add_section_header(_section_title("certifications", lang))
        for cert in certs:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            c_name = _normalize_text(cert.get("name", ""))
            # Clean special characters for ATS
            c_name = c_name.replace("\u2019", "'").replace("\u201c", '"').replace("\u201d", '"')
            c_issuer = _normalize_text(cert.get("issuer", ""))
            c_date = _normalize_text(cert.get("date", ""))
            run = p.add_run(c_name)
            run.font.size = Pt(10.5)
            run.font.name = font.name
            meta = []
            if c_issuer:
                meta.append(c_issuer)
            if c_date:
                meta.append(c_date)
            if meta:
                sep_run = p.add_run(f"  —  {', '.join(meta)}")
                sep_run.font.name = font.name
                sep_run.font.size = Pt(10.5)

    # -- Projects
    projects = cv_data.get("projects", [])

    def _docx_projects():
        if not projects:
            return
        add_section_header(_section_title("projects", lang))
        for proj in projects:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            proj_name = _normalize_text(proj.get("name", ""))
            # Fix "Name- Tech" or "Name -Tech" → "Name - Tech"
            # Only dashes with surrounding whitespace (not compound words like E-Commerce)
            proj_name = re.sub(r"(\w)-\s+(\w)", r"\1 - \2", proj_name)
            proj_name = re.sub(r"(\w)\s+-(\w)", r"\1 - \2", proj_name)
            pn_run = p.add_run(proj_name)
            pn_run.bold = True
            pn_run.font.size = Pt(11)
            pn_run.font.name = font.name

            desc = _normalize_text(proj.get("description", ""))
            if desc:
                dp = doc.add_paragraph()
                dp_run = dp.add_run(desc)
                dp_run.font.size = Pt(10.5)
                dp_run.font.name = font.name

            for bullet in _dedupe_preserve(proj.get("bullets", [])):
                if bullet and bullet.strip():
                    add_bullet(bullet)

    # -- Languages
    languages = cv_data.get("languages", [])

    # -- Interests
    interests = cv_data.get("interests", [])

    # -- Misc
    misc = cv_data.get("misc", [])

    def _docx_languages():
        if not languages:
            return
        add_section_header(_section_title("languages", lang))
        normalized_languages = []
        seen_lang = set()
        for l in languages:
            if isinstance(l, dict):
                key = _normalize_text(l.get("name", "")).lower()
            else:
                key = _normalize_text(str(l)).lower()
            if not key or key in seen_lang:
                continue
            seen_lang.add(key)
            normalized_languages.append(l)
        for l in normalized_languages:
            if isinstance(l, dict):
                name = l.get("name", "")
                writing = l.get("writing") or ""
                listening = l.get("listening") or ""
                speaking = l.get("speaking") or ""
                if writing or listening or speaking:
                    skills = []
                    if writing:
                        skills.append(f"Writing: {writing}")
                    if listening:
                        skills.append(f"Listening: {listening}")
                    if speaking:
                        skills.append(f"Speaking: {speaking}")
                    line = f"{name} \u2013 {', '.join(skills)}"
                else:
                    level = l.get("level", "")
                    line = f"{name} \u2013 {level}" if level else name
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.size = Pt(10.5)
                run.font.name = font.name
            else:
                p = doc.add_paragraph()
                run = p.add_run(str(l))
                run.font.size = Pt(10.5)
                run.font.name = font.name

    # ── Render sections in configurable ATS order ──
    def _docx_interests():
        if not interests:
            return
        add_section_header(_section_title("interests", lang))
        parts = [str(i).strip() for i in interests if str(i or "").strip()]
        import textwrap as _tw

        full_text = ", ".join(parts)
        wrapped = _tw.fill(full_text, width=70)
        for wline in wrapped.split("\n"):
            p = doc.add_paragraph()
            run = p.add_run(wline.strip())
            run.font.size = Pt(10.5)
            run.font.name = font.name

    def _docx_misc():
        if not misc:
            return
        add_section_header(_section_title("misc", lang))
        parts = [str(i).strip() for i in misc if str(i or "").strip()]
        import textwrap as _tw

        full_text = ", ".join(parts)
        wrapped = _tw.fill(full_text, width=70)
        for wline in wrapped.split("\n"):
            p = doc.add_paragraph()
            run = p.add_run(wline.strip())
            run.font.size = Pt(10.5)
            run.font.name = font.name

    _docx_section_renderers = {
        "summary": _docx_summary,
        "experience": _docx_experience,
        "education": _docx_education,
        "skills": _docx_skills,
        "certifications": _docx_certifications,
        "projects": _docx_projects,
        "languages": _docx_languages,
        "interests": _docx_interests,
    }

    for section_name in get_section_order():
        if section_name == "misc":
            continue
        renderer = _docx_section_renderers.get(section_name)
        if renderer:
            renderer()

    # misc always renders last
    _docx_misc()

    # Save to buffer
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# PDF Generation (ATS-optimized, single-column, clean)
# ---------------------------------------------------------------------------


def generate_pdf(cv_data: dict, template: str = "classic", lang: str = "en", font_family: str = "") -> BytesIO:
    """Generate an ATS-friendly PDF. Returns BytesIO buffer."""
    cv_data = _remap_cv_data(cv_data)
    from fpdf import FPDF

    class ATSPDF(FPDF):
        def __init__(self, font_family="Helvetica"):
            super().__init__()
            self._font_family = font_family

        def header(self):
            pass

        def footer(self):
            # Only show page numbers when there are multiple pages
            if self.pages_count > 1:
                self.set_y(-12)
                self.set_font(self._font_family, "", 8)
                self.cell(0, 8, f"{self.page_no()} / {self.pages_count}", align="C")

    # Template-specific PDF styling
    pdf_template_config = {
        "modern": {"font": "Helvetica", "accent_color": (43, 108, 176)},
        "executive": {"font": "Times", "accent_color": (139, 69, 19)},
        "professional": {"font": "Times", "accent_color": (47, 79, 79)},
        "creative": {"font": "Helvetica", "accent_color": (155, 89, 182)},
        "corporate": {"font": "Times", "accent_color": (31, 78, 121)},
        "tech": {"font": "Courier", "accent_color": (0, 122, 204)},
        "consulting": {"font": "Times", "accent_color": (93, 78, 117)},
        "classic": {"font": "Helvetica", "accent_color": (0, 0, 0)},
    }

    pdf_config = pdf_template_config.get(template, pdf_template_config["classic"])
    font_family = pdf_config["font"]
    accent_color = pdf_config["accent_color"]

    pdf = ATSPDF(font_family=font_family)
    try:
        pdf.set_doc_option("core_fonts_encoding", "utf-8")
    except Exception:
        pass

    unicode_font_families = [
        {
            "regular": "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "bold": "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
            "italic": "/usr/share/fonts/truetype/dejavu/DejaVuSans-Oblique.ttf",
        },
        {
            "regular": "/usr/share/fonts/dejavu/DejaVuSans.ttf",
            "bold": "/usr/share/fonts/dejavu/DejaVuSans-Bold.ttf",
            "italic": "/usr/share/fonts/dejavu/DejaVuSans-Oblique.ttf",
        },
        {
            "regular": "C:/Windows/Fonts/DejaVuSans.ttf",
            "bold": "C:/Windows/Fonts/DejaVuSans-Bold.ttf",
            "italic": "C:/Windows/Fonts/DejaVuSans-Oblique.ttf",
        },
        {
            "regular": "C:/Windows/Fonts/arial.ttf",
            "bold": "C:/Windows/Fonts/arialbd.ttf",
            "italic": "C:/Windows/Fonts/ariali.ttf",
        },
        {
            "regular": "C:/Windows/Fonts/calibri.ttf",
            "bold": "C:/Windows/Fonts/calibrib.ttf",
            "italic": "C:/Windows/Fonts/calibrii.ttf",
        },
    ]
    unicode_font_loaded = False
    for family in unicode_font_families:
        regular = family["regular"]
        bold = family["bold"]
        italic = family["italic"]
        if os.path.exists(regular) and os.path.exists(bold):
            try:
                pdf.add_font("UnicodeMain", "", regular)
                pdf.add_font("UnicodeMain", "B", bold)
                if os.path.exists(italic):
                    pdf.add_font("UnicodeMain", "I", italic)
                font_family = "UnicodeMain"
                unicode_font_loaded = True
                break
            except Exception:
                continue
    pdf.set_margins(20, 15, 20)
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    effective_width = pdf.w - pdf.l_margin - pdf.r_margin

    def safe_text(text):
        """Sanitize text for PDF rendering."""
        if not text:
            return ""
        text = _normalize_text(text)
        # Strip markdown bold/italic markers
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        text = re.sub(r"__(.+?)__", r"\1", text)
        text = re.sub(r"(?<!\w)\*(.+?)\*(?!\w)", r"\1", text)
        # Remove stray comment artifacts like /* or */
        text = re.sub(r"/\*|\*/", "", text)
        # Remove remaining leading * or / artifacts
        text = re.sub(r"^[\*/]+\s*", "", text)
        # Fix – • / - • wrap artifacts
        text = text.replace("\u2013 \u2022", "\n\u2022 ")
        text = text.replace("- \u2022", "\n\u2022 ")
        text = text.replace("\u2022 \u2013", "\n\u2022 ")
        text = text.replace("\u2022 -", "\n\u2022 ")
        text = re.sub(r"\s+\u2022", "\n\u2022", text)
        # Normalize smart quotes / ellipsis
        text = text.replace("\u2018", "'")
        text = text.replace("\u2019", "'")
        text = text.replace("\u201c", '"')
        text = text.replace("\u201d", '"')
        text = text.replace("\u2026", "...")
        # Fix "Word(" → "Word ("
        text = re.sub(r"([A-Za-z])\(", r"\1 (", text)
        # Collapse horizontal whitespace only
        text = re.sub(r"[ \t]+", " ", text)
        if unicode_font_loaded:
            return text.strip()
        return text.strip().encode("latin-1", errors="replace").decode("latin-1")

    cv_data = _normalize_contact_fields(cv_data)

    # ── word-aware text helper (never breaks mid-word) ──
    def _break_long_word_cv(word, avail):
        """Break a single long word into chunks that fit within avail."""
        if avail <= 0:
            return [word]
        parts = []
        buf = ""
        for ch in word:
            test = buf + ch
            if pdf.get_string_width(test) > avail and buf:
                parts.append(buf)
                buf = ch
            else:
                buf = test
        if buf:
            parts.append(buf)
        return parts or [word]

    def _write_text(text, style="", size=10, indent=0.0, line_h=5):
        """Write text with word-boundary wrapping. Breaks long words at char level."""
        clean = safe_text(text)
        if not clean:
            return
        # Force no-bold for bullet-like or star-prefixed text
        if clean.startswith(("\u2022", "-", "*")):
            style = ""
        pdf.set_font(font_family, style, size)
        avail = effective_width - indent
        words = re.split(r"\s+", clean)
        lines_buf = []
        cur = ""
        for w in words:
            # If word is wider than available space, break it
            if pdf.get_string_width(w) > avail:
                if cur:
                    lines_buf.append(cur)
                    cur = ""
                lines_buf.extend(_break_long_word_cv(w, avail))
                continue
            cand = f"{cur} {w}".strip() if cur else w
            if pdf.get_string_width(cand) <= avail:
                cur = cand
            else:
                if cur:
                    lines_buf.append(cur)
                cur = w
        if cur:
            lines_buf.append(cur)
        for ln_text in lines_buf:
            if indent:
                pdf.set_x(pdf.l_margin + indent)
            pdf.cell(avail, line_h, ln_text, ln=True)
        # Always reset to normal font after writing
        pdf.set_font(font_family, "", 10)

    # -- Header: name, title, location, contact via shared helper
    name, title_text, location_text, contact_parts = _build_header_data(cv_data)

    if name:
        pdf.set_font(font_family, "B", 16)
        pdf.cell(effective_width, 8, safe_text(name.upper()), align="C", ln=True)
        pdf.ln(1)

    if title_text:
        pdf.set_font(font_family, "", 11)
        pdf.cell(effective_width, 5, safe_text(title_text), align="C", ln=True)
        pdf.ln(0.5)

    if location_text:
        pdf.set_font(font_family, "", 10)
        pdf.cell(effective_width, 5, safe_text(location_text), align="C", ln=True)
        pdf.ln(0.5)

    if contact_parts:
        contact_text = safe_text(" | ".join(contact_parts))
        contact_size = 10
        pdf.set_font(font_family, "", contact_size)
        while pdf.get_string_width(contact_text) > effective_width and contact_size > 7:
            contact_size -= 0.5
            pdf.set_font(font_family, "", contact_size)
        pdf.cell(effective_width, 5, contact_text, align="C", ln=True)
        pdf.ln(2)

    def section_header(title):
        pdf.ln(2)
        pdf.set_font(font_family, "B", 12)

        if accent_color != (0, 0, 0):
            pdf.set_text_color(*accent_color)

        pdf.cell(effective_width, 5, safe_text(title.upper()), ln=True)
        pdf.set_text_color(0, 0, 0)

        if template in ["corporate", "consulting", "executive"]:
            pdf.set_line_width(0.7)
            pdf.set_draw_color(*accent_color)
        elif template in ["creative", "tech"]:
            pdf.set_line_width(0.5)
            pdf.set_draw_color(*accent_color)
        else:
            pdf.set_line_width(0.3)
            pdf.set_draw_color(150, 150, 150)

        pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
        pdf.ln(2)

    def add_bullet(text_content):
        """Render a single bullet point with word-aware wrapping."""
        raw = str(text_content or "").strip()
        if not raw:
            return

        # ── Normalize inline bullets into separate lines ──
        raw = raw.replace(" \u2022 ", "\n\u2022 ")
        raw = raw.replace(" \u2022", "\n\u2022 ")
        raw = re.sub(r"\s+\*\s+", "\n\u2022 ", raw)
        raw = raw.replace("\u2013 \u2022", "\n\u2022")
        raw = raw.replace("- \u2022", "\n\u2022")
        raw = raw.replace("\u2022 \u2013", "\n\u2022")
        raw = raw.replace("\u2022 -", "\n\u2022")

        raw = raw.replace("\u2022", "\n\u2022").replace("\u2023", "\n\u2022")
        parts = []
        for chunk in re.split(r"\n", raw):
            chunk = re.sub(r"^[\u2022\-\*]+\s*", "", chunk).strip()
            if chunk:
                parts.append(chunk)
        if not parts:
            return

        bullet_char = "\u2022" if unicode_font_loaded else "-"
        indent = 10
        bullet_w = 4
        text_w = effective_width - indent - bullet_w

        for part_text in parts:
            clean = safe_text(part_text)
            if not clean:
                continue

            pdf.set_font(font_family, "", 10)
            words = re.split(r"\s+", clean)
            lines_out = []
            current = ""
            for w in words:
                candidate = f"{current} {w}".strip() if current else w
                if pdf.get_string_width(candidate) <= text_w:
                    current = candidate
                else:
                    if current:
                        lines_out.append(current)
                    current = w
            if current:
                lines_out.append(current)

            for li, sline in enumerate(lines_out):
                if li == 0:
                    pdf.set_x(pdf.l_margin + indent)
                    pdf.cell(bullet_w, 5, bullet_char + " ", ln=False)
                    pdf.cell(text_w, 5, sline, ln=True)
                else:
                    pdf.set_x(pdf.l_margin + indent + bullet_w)
                    pdf.cell(text_w, 5, sline, ln=True)
            pdf.ln(0.3)

    # -- Summary
    summary = _normalize_text(cv_data.get("summary", ""))

    def _pdf_summary():
        if not summary:
            return
        section_header(_section_title("summary", lang))
        _write_text(summary)
        pdf.ln(1)

    # -- Experience
    experiences = cv_data.get("experiences", [])

    def _pdf_experience():
        if not experiences:
            return
        section_header(_section_title("experience", lang))
        for exp in experiences:
            # Line 1: Title – Company (bold)
            title_line = _normalize_text(exp.get("title", ""))
            company = _normalize_text(exp.get("company", ""))
            title_line = " \u2013 ".join([x for x in [title_line, company] if x])
            _write_text(title_line, style="B", size=11)

            # Line 2: Date | Location (italic)
            start = _normalize_text(exp.get("start_date", ""))
            end = _normalize_text(exp.get("end_date", ""))
            location = _normalize_text(exp.get("location", ""))
            meta_parts = []
            if start or end:
                date_str = f"{start} - {end}" if start and end else (start or end)
                meta_parts.append(date_str)
            if location:
                meta_parts.append(location)
            if meta_parts:
                pdf.set_font(font_family, "I", 10)
                pdf.cell(effective_width, 5, safe_text(" | ".join(meta_parts)), ln=True)
            pdf.ln(0.5)

            # Bullets
            for bullet in _dedupe_preserve(exp.get("bullets", [])):
                if bullet and bullet.strip():
                    add_bullet(bullet)
            pdf.ln(1)
        pdf.set_font(font_family, "", 10)

    # -- Education
    education = cv_data.get("education", [])

    def _pdf_education():
        if not education:
            return
        section_header(_section_title("education", lang))
        last_gpa = None
        for edu in education:
            degree = _normalize_text(edu.get("degree", ""))
            field = _normalize_text(edu.get("field", ""))
            school = _normalize_text(edu.get("school", ""))

            # Build degree text
            degree_parts = [d.strip() for d in degree.split("|") if d.strip()] if "|" in degree else [degree]

            for i, deg_part in enumerate(degree_parts):
                deg_text = deg_part + (f" in {field}" if field and i == 0 else "")

                # Line 1: Degree (date) — bold
                s = edu.get("start_date", "")
                e = edu.get("end_date", "")
                if s or e:
                    date_str = f"{s} \u2013 {e}" if s and e else (s or e)
                    deg_text = f"{deg_text} ({date_str})" if deg_text.strip() else date_str

                if deg_text.strip():
                    _write_text(deg_text, style="B", size=11)

            # Line 2: School — normal
            if school:
                _write_text(school)

            # Line 3: Location (italic, only if present)
            loc = edu.get("location", "")
            if loc:
                _write_text(loc, style="I")

            # Line 4: GPA (skip if same as previous entry)
            gpa = edu.get("gpa", "")
            if gpa:
                normalized_gpa = _normalize_text(str(gpa)).replace("GPA:", "").replace("gpa:", "").strip()
                if normalized_gpa and normalized_gpa != last_gpa:
                    _write_text(f"GPA: {normalized_gpa}", style="I", size=9)
                    last_gpa = normalized_gpa
            pdf.ln(1.5)
        pdf.set_font(font_family, "", 10)

    # -- Skills
    skills_cat = cv_data.get("skills_categorized", {})

    def _pdf_skills():
        if not skills_cat:
            return
        section_header(_section_title("skills", lang))
        for category, items in skills_cat.items():
            skills_values = _dedupe_preserve(items if isinstance(items, list) else [str(items)])
            skills_values = [_fix_case(v) for v in skills_values]
            skills_text = ", ".join(skills_values)
            cat_label = safe_text(_normalize_text(category)) + ": "
            full_line = cat_label + safe_text(skills_text)

            # Word-aware wrapping: split into lines that fit effective_width
            pdf.set_font(font_family, "", 10)
            words = full_line.split(" ")
            lines_out = []
            current = ""
            for w in words:
                candidate = f"{current} {w}".strip() if current else w
                if pdf.get_string_width(candidate) <= effective_width:
                    current = candidate
                else:
                    if current:
                        lines_out.append(current)
                    current = w
            if current:
                lines_out.append(current)

            for li, sline in enumerate(lines_out):
                if li == 0:
                    # First line: bold category, normal values
                    pdf.set_font(font_family, "B", 10)
                    cat_w = pdf.get_string_width(cat_label) + 1
                    pdf.cell(cat_w, 5, cat_label, ln=False)
                    pdf.set_font(font_family, "", 10)
                    remainder = sline[len(cat_label) :].strip()
                    pdf.cell(effective_width - cat_w, 5, remainder, ln=True)
                else:
                    pdf.set_font(font_family, "", 10)
                    indent = pdf.get_string_width("    ")
                    pdf.set_x(pdf.l_margin + indent)
                    pdf.cell(effective_width - indent, 5, sline, ln=True)
        pdf.set_font(font_family, "", 10)

    # -- Certifications
    certs = cv_data.get("certifications", [])

    def _pdf_certifications():
        if not certs:
            return
        section_header(_section_title("certifications", lang))
        for cert in certs:
            cert_name = _normalize_text(cert.get("name", ""))
            # Clean special characters for ATS
            cert_name = cert_name.replace("\u2019", "'").replace("\u201c", '"').replace("\u201d", '"')
            cert_name = safe_text(cert_name)
            meta = []
            if cert.get("issuer"):
                meta.append(cert["issuer"])
            if cert.get("date"):
                meta.append(cert["date"])
            if meta:
                line = f"{cert_name} - {', '.join(meta)}"
            else:
                line = cert_name
            _write_text(line)
            pdf.ln(0.5)

    # -- Projects
    projects = cv_data.get("projects", [])

    def _pdf_projects():
        if not projects:
            return
        section_header(_section_title("projects", lang))
        for proj in projects:
            name = _normalize_text(proj.get("name", ""))
            desc = _normalize_text(proj.get("description", ""))

            # Fix "Name- Tech" or "Name -Tech" → "Name - Tech"
            # Only fix dashes with surrounding whitespace (not compound words like E-Commerce)
            name = re.sub(r"(\w)-\s+(\w)", r"\1 - \2", name)  # "Name- Tech"
            name = re.sub(r"(\w)\s+-(\w)", r"\1 - \2", name)  # "Name -Tech"

            _write_text(name, style="B", size=11)

            if desc:
                _write_text(desc)
            for bullet in _dedupe_preserve(proj.get("bullets", [])):
                if bullet and bullet.strip():
                    add_bullet(bullet)
            pdf.ln(1)
        pdf.set_font(font_family, "", 10)

    # -- Languages
    languages = cv_data.get("languages", [])

    # -- Interests
    interests = cv_data.get("interests", [])

    # -- Misc
    misc = cv_data.get("misc", [])

    def _pdf_languages():
        if not languages:
            return
        section_header(_section_title("languages", lang))
        for l in languages:
            if isinstance(l, dict):
                name = l.get("name", "")
                writing = l.get("writing") or ""
                listening = l.get("listening") or ""
                speaking = l.get("speaking") or ""
                if writing or listening or speaking:
                    skills = []
                    if writing:
                        skills.append(f"Writing: {writing}")
                    if listening:
                        skills.append(f"Listening: {listening}")
                    if speaking:
                        skills.append(f"Speaking: {speaking}")
                    _write_text(f"{name} - {', '.join(skills)}")
                else:
                    level = l.get("level", "")
                    _write_text(f"{name} - {level}" if level else name)
            else:
                _write_text(str(l))

    def _pdf_interests():
        if not interests:
            return
        section_header(_section_title("interests", lang))
        parts = [str(i).strip() for i in interests if str(i or "").strip()]
        _write_text(", ".join(parts))

    # ── Render sections in configurable ATS order ──
    def _pdf_misc():
        if not misc:
            return
        section_header(_section_title("misc", lang))
        parts = [str(i).strip() for i in misc if str(i or "").strip()]
        _write_text(", ".join(parts))

    _pdf_section_renderers = {
        "summary": _pdf_summary,
        "experience": _pdf_experience,
        "education": _pdf_education,
        "skills": _pdf_skills,
        "certifications": _pdf_certifications,
        "projects": _pdf_projects,
        "languages": _pdf_languages,
        "interests": _pdf_interests,
    }

    for section_name in get_section_order():
        if section_name == "misc":
            continue
        renderer = _pdf_section_renderers.get(section_name)
        if renderer:
            renderer()

    # misc always renders last
    _pdf_misc()

    pdf_output = pdf.output(dest="S")
    pdf_bytes = pdf_output.encode("latin-1") if isinstance(pdf_output, str) else bytes(pdf_output)

    buf = BytesIO(pdf_bytes)
    buf.seek(0)
    return buf


def generate_typst(cv_data: dict, template: str = "classic", font_family: str = "") -> BytesIO:
    """Generate a Typst source document buffer.

    This keeps rendering deterministic and avoids runtime dependency on a
    Typst binary in API workers. Clients can compile `.typ` to PDF.
    """

    def esc(value: str) -> str:
        return str(value or "").replace("\\", "\\\\").replace('"', '\\"')

    # Resolve font: user override > template default > Liberation Sans
    from renderers.theme import ALLOWED_FONTS

    typst_font = "Liberation Sans"
    if font_family and font_family in ALLOWED_FONTS:
        typst_font = font_family

    full_name = esc(cv_data.get("full_name", ""))
    email = esc(cv_data.get("email", ""))
    phone = esc(cv_data.get("phone", ""))
    location = esc(cv_data.get("location", ""))
    linkedin = esc(cv_data.get("linkedin", ""))
    summary = esc(cv_data.get("summary", ""))

    lines: list[str] = [
        "#set page(margin: 1.5cm)",
        f'#set text(font: "{esc(typst_font)}", size: 10pt)',
        f'#let templateName = "{esc(template)}"',
        "",
        f"= {full_name or 'Curriculum Vitae'}",
        f"{email}  {phone}  {location}  {linkedin}".strip(),
        "",
    ]

    if summary:
        lines += ["== Professional Summary", summary, ""]

    experiences = cv_data.get("experiences") or []
    if experiences:
        lines.append("== Experience")
        for exp in experiences:
            if not isinstance(exp, dict):
                continue
            title = esc(exp.get("title", ""))
            company = esc(exp.get("company", ""))
            start = esc(exp.get("start_date", ""))
            end = esc(exp.get("end_date", ""))
            header = " - ".join(part for part in [title, company] if part)
            date_span = " - ".join(part for part in [start, end] if part)
            if header:
                lines.append(f"=== {header}")
            if date_span:
                lines.append(date_span)
            for bullet in exp.get("bullets", []) or []:
                if str(bullet or "").strip():
                    lines.append(f"- {esc(str(bullet).strip())}")
            lines.append("")

    education = cv_data.get("education") or []
    if education:
        lines.append("== Education")
        for item in education:
            if isinstance(item, dict):
                degree = esc(item.get("degree", ""))
                school = esc(item.get("school", ""))
                if degree or school:
                    lines.append(" - ".join(part for part in [degree, school] if part))
            else:
                lines.append(esc(str(item)))
        lines.append("")

    skills = cv_data.get("skills") or []
    if skills:
        lines += ["== Skills", ", ".join(esc(str(s)) for s in skills if str(s).strip()), ""]

    projects = cv_data.get("projects") or []
    if projects:
        lines.append("== Projects")
        for p in projects:
            if isinstance(p, dict):
                name = esc(p.get("name", ""))
                desc = esc(p.get("description", ""))
                if name:
                    lines.append(f"=== {name}")
                if desc:
                    lines.append(desc)
                for bullet in p.get("bullets", []) or []:
                    if str(bullet or "").strip():
                        lines.append(f"- {esc(str(bullet).strip())}")
                lines.append("")

    languages = cv_data.get("languages") or []
    if languages:
        lines += ["== Languages", ", ".join(esc(str(l)) for l in languages if str(l).strip()), ""]

    content = "\n".join(lines).strip() + "\n"
    buf = BytesIO(content.encode("utf-8"))
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

# All CV templates are ATS-compliant (Application Tracking System)
# They use standard fonts, clean formatting, and proper structure
# to ensure compatibility with automated resume scanning systems
TEMPLATES = {
    "free": ["classic"],
    "pro": ["classic", "modern", "executive", "professional", "creative"],
    "enterprise": ["classic", "modern", "executive", "professional", "creative", "corporate", "tech", "consulting"],
    "admin": ["classic", "modern", "executive", "professional", "creative", "corporate", "tech", "consulting"],
}


def get_available_templates(plan: str) -> list:
    return TEMPLATES.get(plan, TEMPLATES["free"])


def _is_truthy_env(name: str, default: str = "0") -> bool:
    value = str(os.getenv(name, default) or default).strip().lower()
    return value in ("1", "true", "yes", "on")


def _payload_has_missing_core_sections(payload: dict) -> bool:
    if not isinstance(payload, dict):
        return True
    summary_missing = not str(payload.get("summary") or "").strip()
    experiences_missing = not isinstance(payload.get("experiences"), list) or len(payload.get("experiences") or []) == 0
    education_missing = not isinstance(payload.get("education"), list) or len(payload.get("education") or []) == 0

    skills_categorized = payload.get("skills_categorized") or {}
    plain_skills = payload.get("skills") or []
    skills_missing = (not isinstance(skills_categorized, dict) or len(skills_categorized) == 0) and (
        not isinstance(plain_skills, list) or len(plain_skills) == 0
    )

    return summary_missing or experiences_missing or education_missing or skills_missing


def _payload_to_ats_text(payload: dict) -> str:
    data = dict(payload or {})
    lines = []

    if str(data.get("full_name") or "").strip():
        lines.append(str(data.get("full_name")).strip())

    contacts = [
        str(data.get("email") or "").strip(),
        str(data.get("phone") or "").strip(),
        str(data.get("location") or "").strip(),
    ]
    contacts = [c for c in contacts if c]
    if contacts:
        lines.append(" | ".join(contacts))

    summary = str(data.get("summary") or "").strip()
    if summary:
        lines.extend(["", "PROFESSIONAL SUMMARY", summary])

    experiences = data.get("experiences") or []
    if isinstance(experiences, list) and experiences:
        lines.extend(["", "EXPERIENCE"])
        for exp in experiences:
            if not isinstance(exp, dict):
                continue
            title = str(exp.get("title") or "").strip()
            company = str(exp.get("company") or "").strip()
            header = " - ".join([p for p in [title, company] if p])
            if header:
                lines.append(header)
            dates = " - ".join(
                [
                    str(exp.get("start_date") or "").strip(),
                    str(exp.get("end_date") or "").strip(),
                ]
            ).strip(" -")
            if dates:
                lines.append(dates)
            for bullet in exp.get("bullets") or []:
                text = str(bullet or "").strip()
                if text:
                    lines.append(f"- {text}")

    projects = data.get("projects") or []
    if isinstance(projects, list) and projects:
        lines.extend(["", "PROJECTS"])
        for proj in projects:
            if not isinstance(proj, dict):
                continue
            name = str(proj.get("name") or "").strip()
            desc = str(proj.get("description") or "").strip()
            if name:
                lines.append(name)
            if desc:
                lines.append(desc)
            for bullet in proj.get("bullets") or []:
                text = str(bullet or "").strip()
                if text:
                    lines.append(f"- {text}")

    education = data.get("education") or []
    if isinstance(education, list) and education:
        lines.extend(["", "EDUCATION"])
        for ed in education:
            if not isinstance(ed, dict):
                continue
            degree = str(ed.get("degree") or "").strip()
            school = str(ed.get("school") or "").strip()
            row = " - ".join([p for p in [degree, school] if p])
            if row:
                lines.append(row)

    skills_map = data.get("skills_categorized") or {}
    if isinstance(skills_map, dict) and skills_map:
        lines.extend(["", "SKILLS"])
        for category, values in skills_map.items():
            values = [str(v).strip() for v in (values or []) if str(v).strip()]
            if not values:
                continue
            lines.append(str(category).strip())
            for value in values:
                lines.append(f"- {value}")
    else:
        plain_skills = [str(v).strip() for v in (data.get("skills") or []) if str(v).strip()]
        if plain_skills:
            lines.extend(["", "SKILLS", ", ".join(plain_skills)])

    return "\n".join(lines).strip()


def build_cv(
    cv_data: dict,
    job_description: str,
    template: str = "classic",
    output_format: str = "docx",
    lang: str = "en",
    plan: str = "free",
    font_family: str = "",
) -> dict:
    """
    Main entry: enhance CV data with AI, generate document.
    Returns dict with 'buffer' (BytesIO), 'filename', 'content_type'.
    """
    _build_t0 = time.perf_counter()

    # ── Config snapshot: log full runtime config at request entry ──
    _structured_log(
        logger,
        logging.INFO,
        "cv_parse",
        **_config_snapshot(),
    )

    # ── Global circuit breaker ──
    acquired = _global_parse_semaphore.acquire(blocking=False)
    if not acquired:
        logger.warning(
            "Global parse limit reached (%d); rejecting request",
            _GLOBAL_PARSE_LIMIT,
        )
        raise RuntimeError(f"Server overloaded: {_GLOBAL_PARSE_LIMIT} concurrent parses in progress")
    try:
        return _build_cv_inner(
            cv_data,
            job_description,
            template,
            output_format,
            lang,
            plan,
            _build_t0,
            font_family=font_family,
        )
    except Exception:
        _record_parse_error()
        raise
    finally:
        _global_parse_semaphore.release()


def _build_cv_inner(
    cv_data: dict,
    job_description: str,
    template: str,
    output_format: str,
    lang: str,
    plan: str,
    _build_t0: float,
    font_family: str = "",
) -> dict:
    # ── Safe / degraded mode: skip classifier + AI, go straight to raw render ──
    if _is_safe_mode():
        logger.warning("SAFE_MODE active — bypassing classifier and AI, rendering raw text")
        cv_data = dict(cv_data or {})
        name = str(cv_data.get("full_name", "CV"))
        cv_model = CVModel(full_name=name)
        try:
            buf = _render_raw_text_fallback(cv_data, cv_model)
        except Exception:
            buf = _emergency_pdf_fallback()
        safe_name = re.sub(r"[^a-zA-Z0-9_\- ]", "", name).strip().replace(" ", "_") or "CV"
        return {
            "buffer": buf,
            "filename": f"{safe_name}_CV.pdf",
            "content_type": "application/pdf",
            "enhanced_data": cv_model.model_dump(),
            "cache_hit": False,
            "ai_final_review_applied": False,
        }

    # Validate template against plan-allowed list; fall back to ATS default if invalid.
    allowed = get_available_templates(plan)
    if template not in allowed:
        template = ATS_DEFAULT_TEMPLATE

    # Canonicalize input keys so builder never depends on singular/plural variants.
    cv_data = dict(cv_data or {})
    if "experiences" not in cv_data and "experience" in cv_data:
        cv_data["experiences"] = cv_data.get("experience") or []
    if "projects" not in cv_data and "project" in cv_data:
        cv_data["projects"] = cv_data.get("project") or []
    if "languages" not in cv_data and "language" in cv_data:
        val = cv_data.get("language") or []
        # "language" can be a detection code like "tr" — only promote to "languages" if it's a list
        if isinstance(val, list):
            cv_data["languages"] = val
    if "skills" not in cv_data and "skill" in cv_data:
        cv_data["skills"] = cv_data.get("skill") or []

    # Parse -> compile -> ATS gate -> AI final review -> render
    enhanced = dict(cv_data or {})
    enhanced = _normalize_contact_fields(enhanced)

    try:
        cv_model = compile_cv_model(enhanced)
    except Exception:
        logger.exception("compile_cv_model failed; building minimal model from raw data")
        cv_model = CVModel(full_name=str(cv_data.get("full_name", "CV")))

    # Prefer the auto-detected language from the CV over the caller-provided lang.
    model_lang = getattr(cv_model, "language", "") or ""
    if model_lang and model_lang != "en":
        lang = model_lang
    elif not lang:
        lang = model_lang or "en"

    # ── Slow-CV guard: if compilation already took too long, skip AI review ──
    _elapsed_so_far = time.perf_counter() - _build_t0
    _skip_ai_review = False
    if _elapsed_so_far > _SLOW_ABORT_SECONDS:
        logger.warning(
            "build_cv SLOW ABORT: %.2fs elapsed after compile (limit %.1fs), "
            "skipping AI review and rendering immediately",
            _elapsed_so_far,
            _SLOW_ABORT_SECONDS,
        )
        _skip_ai_review = True
    elif _elapsed_so_far > _SLOW_WARN_SECONDS:
        logger.warning(
            "build_cv SLOW: %.2fs elapsed after compile (warn threshold %.1fs)",
            _elapsed_so_far,
            _SLOW_WARN_SECONDS,
        )

    final_review_enabled = _is_truthy_env("AI_FINAL_REVIEW", "1")
    final_review_only_missing = _is_truthy_env("AI_FINAL_REVIEW_ONLY_MISSING", "0")
    final_review_require_missing = _is_truthy_env("AI_FINAL_REVIEW_REQUIRE_MISSING", "1")
    final_review_threshold = float(os.getenv("AI_FINAL_REVIEW_ATS_THRESHOLD", "60") or "60")

    ai_final_review_applied = False

    if final_review_enabled and ENABLE_AI_REVIEW and not _skip_ai_review:
        try:
            review_input = cv_model.model_dump(mode="json")
            should_review = False

            ats_text = _payload_to_ats_text(review_input)
            if "\n\n\n" not in ats_text:
                ats_text = ats_text.replace("  ", "\n")
            ats_snapshot = analyze_cv(ats_text, job_description or "", lang=lang)
            ats_score = float((ats_snapshot or {}).get("overall_score", 0) or 0)
            has_missing_sections = _payload_has_missing_core_sections(review_input)

            if ats_score < final_review_threshold:
                should_review = True

            # Default behavior: review only when score is low AND core sections are missing.
            if final_review_require_missing:
                should_review = should_review and has_missing_sections

            # Optional strict mode retained for backward compatibility.
            if final_review_only_missing:
                should_review = should_review and has_missing_sections

            if should_review:
                reviewed_payload = rewrite_service.ai_review_cv_payload(
                    review_input,
                    job_description=job_description,
                    lang=lang,
                )
                reviewed_payload = _normalize_contact_fields(reviewed_payload)
                cv_model = compile_cv_model(reviewed_payload)
                ai_final_review_applied = True
        except Exception:
            logger.exception("AI final review step failed; continuing with pre-review payload")

    cache_key = make_cache_key(
        {
            "template": template,
            "output_format": output_format,
            "plan": plan,
            "model": cv_model.model_dump(mode="json"),
        }
    )

    cached = get_cached(cache_key)
    if cached:
        buf = BytesIO(cached["bytes"])
        filename = cached["filename"]
        content_type = cached["content_type"]
        return {
            "buffer": buf,
            "filename": filename,
            "content_type": content_type,
            "enhanced_data": cv_model.model_dump(),
            "cache_hit": True,
            "ai_final_review_applied": ai_final_review_applied,
        }

    full_name = cv_model.full_name or cv_data.get("full_name", "CV")
    safe_name = re.sub(r"[^a-zA-Z0-9_\- ]", "", full_name).strip().replace(" ", "_")
    if not safe_name:
        safe_name = "CV"

    # ── Pre-render sanity check ──
    if ENABLE_SANITIZER:
        try:
            _pre_render_sanity_check(cv_model)
        except Exception:
            logger.exception("_pre_render_sanity_check failed; continuing with unchecked model")

    # ── Fallback: never return empty CV ──
    if ENABLE_FALLBACK and _needs_fallback_render(cv_model):
        logger.info("build_cv: fallback triggered by _needs_fallback_render")
        cv_model = _build_fallback_model(cv_data, cv_model)

    # ── Final layout normalization ──
    _normalize_layout_model(cv_model)

    # ── Schema integrity: ensure primary sections exist ──
    try:
        _model_integrity_check(cv_model)
    except Exception:
        logger.exception("_model_integrity_check failed; continuing")

    # ── Sanity score: last-resort fallback if all primary sections empty ──
    _sanity = _model_sanity_score(cv_model)
    if _sanity == 0:
        logger.warning("build_cv: model_sanity_score=0, triggering fallback model")
        cv_model = _build_fallback_model(cv_data, cv_model)

    logger.info(
        "build_cv: sanity=%d | exp=%d edu=%d | format=%s",
        _sanity,
        len(cv_model.experiences),
        len(cv_model.education),
        output_format,
    )
    _structured_log(
        logger,
        logging.INFO,
        "build_cv_summary",
        sanity=_sanity,
        experiences=len(cv_model.experiences),
        education=len(cv_model.education),
        format=output_format,
        safe_mode=_is_safe_mode(),
        ai_review=ai_final_review_applied,
        latency=round(time.perf_counter() - _build_t0, 3),
        version=PARSER_VERSION,
        build_id=BUILD_ID,
        git_sha=GIT_SHA,
    )

    data = _reorder_data_dict(cv_model.model_dump())

    try:
        if output_format == "pdf":
            if USE_SCHEMA_PIPELINE:
                buf = _render_pdf_blocks(cv_model, template, font_override=font_family)
            else:
                buf = generate_pdf(data, template, lang=lang, font_family=font_family)
            extension = "pdf"
            content_type = "application/pdf"
        elif output_format == "docx":
            buf = generate_docx(data, template, lang=lang, font_family=font_family)
            extension = "docx"
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif output_format == "typst":
            buf = generate_typst(data, template, font_family=font_family)
            extension = "typ"
            content_type = "text/plain"
        elif output_format == "html":
            from renderers import render as _render_all

            r = _render_all(cv_model, template, "html", font_override=font_family)
            buf = r["buffer"]
            extension = "html"
            content_type = "text/html; charset=utf-8"
        else:
            raise ValueError(f"Unsupported format: {output_format}")
    except Exception:
        logger.exception("Render failed for %s; falling back to raw text PDF", output_format)
        try:
            buf = _render_raw_text_fallback(cv_data, cv_model)
        except Exception:
            logger.exception("Raw text fallback also failed; returning minimal PDF")
            buf = _emergency_pdf_fallback()
        extension = "pdf"
        content_type = "application/pdf"

    filename = f"{safe_name}_CV.{extension}"

    rendered_bytes = buf.read()
    buf = BytesIO(rendered_bytes)
    set_cached(
        cache_key,
        {
            "bytes": rendered_bytes,
            "filename": filename,
            "content_type": content_type,
        },
    )

    return {
        "buffer": buf,
        "filename": filename,
        "content_type": content_type,
        "enhanced_data": cv_model.model_dump(),
        "cache_hit": False,
        "ai_final_review_applied": ai_final_review_applied,
    }


def _render_raw_text_fallback(cv_data: dict, cv_model: CVModel) -> BytesIO:
    """Last-resort renderer: dump all available text into a minimal PDF."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Helvetica", size=11)

    name = cv_model.full_name or cv_data.get("full_name", "CV")
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, name, ln=True)
    pdf.set_font("Helvetica", size=10)

    # Collect all text from cv_data
    for key, value in cv_data.items():
        if key.startswith("_") or not value:
            continue
        if isinstance(value, str) and value.strip():
            pdf.multi_cell(0, 5, value.strip())
            pdf.ln(2)
        elif isinstance(value, list):
            for item in value:
                if isinstance(item, str) and item.strip():
                    pdf.multi_cell(0, 5, item.strip())
                elif isinstance(item, dict):
                    for v in item.values():
                        if isinstance(v, str) and v.strip():
                            pdf.multi_cell(0, 5, v.strip())

    buf = BytesIO()
    buf.write(pdf.output())
    buf.seek(0)
    return buf


def _emergency_pdf_fallback() -> BytesIO:
    """Absolute last-resort: return a tiny valid PDF with an error message.

    This must never raise.  Uses fpdf with only ASCII text to guarantee
    that even encoding issues won't break it.
    """
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.cell(0, 10, "CV document could not be generated.", ln=True)
    pdf.cell(0, 10, "Please try again or contact support.", ln=True)
    buf = BytesIO()
    buf.write(pdf.output())
    buf.seek(0)
    return buf


def compile_cv_model(cv_data: dict | CVModel) -> CVModel:
    if isinstance(cv_data, CVModel):
        model = cv_data
    else:
        model = CVModel.from_mapping(cv_data)

    # -------- basic --------
    model.full_name = _normalize_text(model.full_name).strip()
    model.email = _normalize_text(model.email).strip()
    model.phone = _normalize_text(model.phone).strip()
    model.location = _normalize_text(model.location).strip()
    model.summary = _normalize_text(model.summary).strip()
    model.linkedin = _normalize_text(getattr(model, "linkedin", "")).strip()

    # -------- experiences --------
    for exp in model.experiences:
        exp.title = _normalize_text(exp.title).strip()
        exp.company = _normalize_text(exp.company).strip()
        exp.location = _normalize_text(exp.location).strip()
        exp.start_date = _normalize_text(exp.start_date).strip()
        exp.end_date = _normalize_text(exp.end_date).strip()
        exp.bullets = _dedupe_preserve([_normalize_text(v).strip() for v in exp.bullets])

    # -------- education --------
    for edu in model.education:
        edu.degree = _normalize_text(edu.degree).strip()
        edu.school = _normalize_text(edu.school).strip()
        edu.start_date = _normalize_text(edu.start_date).strip()
        edu.end_date = _normalize_text(edu.end_date).strip()
        edu.gpa = _normalize_text(edu.gpa).strip()
        edu.field = _normalize_text(edu.field).strip()
        edu.location = _normalize_text(edu.location).strip()

    # normalize multiple degrees inside same entry (double major fix)
    extra_edu = []
    for edu in model.education:
        if "|" in edu.degree:
            parts = [p.strip() for p in edu.degree.split("|")]
            edu.degree = parts[0]
            for extra_deg in parts[1:]:
                extra_edu.append(
                    Education(
                        degree=extra_deg,
                        school=edu.school,
                        start_date=edu.start_date,
                        end_date=edu.end_date,
                        gpa="",
                        field=edu.field,
                        location=edu.location,
                    )
                )
    model.education.extend(extra_edu)

    # -------- projects --------
    for p in model.projects:
        p.name = _normalize_text(p.name).strip()
        p.description = _normalize_text(p.description).strip()
        p.bullets = _dedupe_preserve([_normalize_text(b).strip() for b in p.bullets if _normalize_text(b).strip()])
        # GUARD: if project has description but no bullets, promote description to bullet
        if p.description and not p.bullets:
            p.bullets = [p.description]
            p.description = ""

    # -------- certifications --------
    for c in model.certifications:
        c.name = _normalize_text(c.name).strip()
        c.issuer = _normalize_text(c.issuer).strip()
        c.date = _normalize_text(c.date).strip()

    # -------- skills --------
    cleaned_skills = _dedupe_preserve([_normalize_text(s).strip() for s in model.skills])
    model.skills = cleaned_skills

    if model.skills_categorized:
        fixed = {}
        for k, v in model.skills_categorized.items():
            key = _normalize_text(k).strip()
            vals = _dedupe_preserve([_normalize_text(x).strip() for x in v])
            fixed[key] = vals
        model.skills_categorized = fixed
    elif cleaned_skills and not model.skills_categorized:
        model.skills_categorized = {"Technical Skills": cleaned_skills}

    # -------- limit skills per category --------
    for k in model.skills_categorized:
        model.skills_categorized[k] = model.skills_categorized[k][:25]

    # -------- languages --------
    # GUARD: languages must never be lost after parsing
    if hasattr(model, "languages") and model.languages:
        original_count = len(model.languages)
        model.languages = _dedupe_preserve(
            [_normalize_text(x).strip() for x in model.languages if _normalize_text(x).strip()]
        )
        # Safety: if dedup removed everything, restore originals
        if not model.languages and original_count > 0:
            model.languages = (
                _dedupe_preserve(
                    [
                        str(x).strip()
                        for x in (cv_data if isinstance(cv_data, dict) else {}).get("languages", [])
                        if str(x).strip()
                    ]
                )
                or model.languages
            )

    return model


# ── Pre-render sanity check ──────────────────────────────────────────────


def _pre_render_sanity_check(model: CVModel) -> None:
    """Final in-place cleanup before document rendering.

    * Remove invalid language entries (single char, pure numbers, URLs).
    * Remove empty contact fields.
    * Drop empty misc.
    * Fix education entries that have no school AND no degree.
    * Fix skills that are really dates or URLs.
    * Section size guard on experience bullets.
    """
    # Languages
    if hasattr(model, "languages") and model.languages:
        model.languages = [
            lang
            for lang in model.languages
            if lang
            and len(lang) > 1
            and not re.match(r"^[\d\W]+$", lang)
            and "@" not in lang
            and not re.match(r"https?://", lang, re.I)
        ]

    # Contact
    if model.email and not re.search(r"@.+\.", model.email):
        model.email = ""
    if model.phone and len(re.sub(r"\D", "", model.phone)) < 8:
        model.phone = ""
    if hasattr(model, "linkedin"):
        if model.linkedin and not re.search(
            r"(?:https?://|www\.|(?:[a-z0-9-]+\.)+[a-z]{2,}(?:/|$))",
            model.linkedin,
            re.I,
        ):
            model.linkedin = ""

    # Education: drop empty entries
    model.education = [edu for edu in model.education if edu.degree or edu.school]

    # Skills: remove date/URL items
    _date_like = re.compile(r"^\d{4}\s*[-–]\s*\d{4}$|^\d{1,2}/\d{4}$")
    model.skills = [s for s in model.skills if s and not _date_like.match(s) and not re.match(r"https?://", s, re.I)]
    if model.skills_categorized:
        for cat in list(model.skills_categorized):
            cleaned = [
                s
                for s in model.skills_categorized[cat]
                if s and not _date_like.match(s) and not re.match(r"https?://", s, re.I)
            ]
            if cleaned:
                model.skills_categorized[cat] = cleaned
            else:
                del model.skills_categorized[cat]

    # Misc: drop empties
    if hasattr(model, "misc"):
        model.misc = [m for m in (model.misc or []) if m]

    # Experience bullet size guard: if any single experience has >20 bullets,
    # truncate to a safe limit (likely parser noise).
    for exp in model.experiences:
        if len(exp.bullets) > 20:
            exp.bullets = exp.bullets[:20]

    # Cross-section: move experience entries that look like education
    _cross_section_fixup_model(model)

    # Document-level validation: final structural consistency
    _document_level_validation_model(model)

    # Anomaly detection: final abnormal-structure fix
    _anomaly_detection_model(model)


# ── Cross-section fixup (CVModel) ──────────────────────────────────

_DEGREE_MODEL_RE = re.compile(
    r"\b(?:b\.?s\.?c?|m\.?s\.?c?|b\.?a|m\.?a|ph\.?d|m\.?b\.?a"
    r"|bachelor|master|diploma|associate|degree)\b",
    re.I,
)
_INSTITUTION_MODEL_RE = re.compile(
    r"\b(?:university|institute|college|school|faculty|academy)\b",
    re.I,
)


def _cross_section_fixup_model(model: CVModel) -> None:
    """Final cross-section correction on the CVModel before rendering."""
    # Experience → Education: entry with degree/university, no bullets, has date
    kept: list = []
    for exp in model.experiences:
        text = f"{exp.title} {exp.company} {exp.location} {exp.start_date} {exp.end_date}"
        has_degree = bool(_DEGREE_MODEL_RE.search(text))
        has_institution = bool(_INSTITUTION_MODEL_RE.search(text))
        has_date = bool(re.search(r"\b(?:19|20)\d{2}\b", text))
        if (has_degree or has_institution) and has_date and not exp.bullets:
            model.education.append(
                Education(
                    degree=exp.title,
                    school=exp.company,
                    location=exp.location,
                    start_date=exp.start_date,
                    end_date=exp.end_date,
                )
            )
        else:
            kept.append(exp)
    model.experiences = kept

    # Deduplicate education (same school + degree)
    seen_edu: set[str] = set()
    deduped_edu: list = []
    for edu in model.education:
        key = f"{edu.school.lower().strip()}|{edu.degree.lower().strip()}"
        if key not in seen_edu:
            seen_edu.add(key)
            deduped_edu.append(edu)
    model.education = deduped_edu

    # Language validation: reject tech names, URLs, pure numbers
    if hasattr(model, "languages") and model.languages:
        _tech_re = re.compile(
            r"\b(?:python|javascript|typescript|react|angular|docker|kubernetes"
            r"|sql|html|css|git|aws|azure|node\.?js)\b",
            re.I,
        )
        model.languages = [
            lang
            for lang in model.languages
            if lang
            and len(lang.strip()) > 1
            and not _tech_re.search(lang)
            and not re.match(r"https?://", lang, re.I)
            and not re.match(r"^[\d\W]+$", lang.strip())
            and "@" not in lang
        ]


# ── Document-level validation (CVModel) ──────────────────────────────

_YEAR_MODEL_RE = re.compile(r"\b(?:19|20)\d{2}\b")
_MAX_MODEL_ENTRIES = 50
_MAX_BULLETS_PER_ENTRY_MODEL = 20


def _document_level_validation_model(model: CVModel) -> None:
    """Final document-level consistency on CVModel before rendering.

    Only *moves* existing data — never creates new content.
    """

    # ── 1. Section-size sanity ──
    if len(model.experiences) > _MAX_MODEL_ENTRIES:
        keep: list = []
        for exp in model.experiences:
            text = f"{exp.title} {exp.company} {exp.location} {exp.start_date} {exp.end_date}"
            if (_DEGREE_MODEL_RE.search(text) or _INSTITUTION_MODEL_RE.search(text)) and not exp.bullets:
                model.education.append(
                    Education(
                        degree=exp.title,
                        school=exp.company,
                        location=exp.location,
                        start_date=exp.start_date,
                        end_date=exp.end_date,
                    )
                )
            else:
                keep.append(exp)
        model.experiences = keep

    if len(model.education) > _MAX_MODEL_ENTRIES:
        model.education = model.education[:_MAX_MODEL_ENTRIES]
    if len(model.skills) > 100:
        model.skills = model.skills[:100]
    if len(model.projects) > _MAX_MODEL_ENTRIES:
        model.projects = model.projects[:_MAX_MODEL_ENTRIES]

    # Bullet cap per experience/project entry
    for exp in model.experiences:
        if len(exp.bullets) > _MAX_BULLETS_PER_ENTRY_MODEL:
            exp.bullets = exp.bullets[:_MAX_BULLETS_PER_ENTRY_MODEL]
    for proj in model.projects:
        if len(proj.bullets) > _MAX_BULLETS_PER_ENTRY_MODEL:
            proj.bullets = proj.bullets[:_MAX_BULLETS_PER_ENTRY_MODEL]

    # ── 2. Primary-section check ──
    _order = [
        ("summary", bool(model.summary)),
        ("experiences", bool(model.experiences)),
        ("education", bool(model.education)),
        ("skills", bool(model.skills or model.skills_categorized)),
        ("languages", bool(model.languages)),
        ("misc", bool(getattr(model, "misc", None))),
    ]
    first = None
    for name, has in _order:
        if has:
            first = name
            break

    if first == "misc" and hasattr(model, "misc") and model.misc:
        promoted: list[str] = []
        kept_misc: list[str] = []
        for item in model.misc:
            if len(item.split()) >= 8:
                promoted.append(item)
            else:
                kept_misc.append(item)
        if promoted:
            extra = " ".join(promoted)
            model.summary = f"{model.summary} {extra}".strip() if model.summary else extra
            model.misc = kept_misc

    # ── 3. Education existence check ──
    if not model.education:
        remaining: list = []
        for exp in model.experiences:
            text = f"{exp.title} {exp.company} {exp.start_date} {exp.end_date}"
            if _DEGREE_MODEL_RE.search(text) and _YEAR_MODEL_RE.search(text) and not exp.bullets:
                model.education.append(
                    Education(
                        degree=exp.title,
                        school=exp.company,
                        location=exp.location,
                        start_date=exp.start_date,
                        end_date=exp.end_date,
                    )
                )
            else:
                remaining.append(exp)
        model.experiences = remaining

    # ── 4. Experience existence check ──
    if not model.experiences:
        remaining_edu: list = []
        for edu in model.education:
            text = f"{edu.degree} {edu.school} {edu.start_date} {edu.end_date}"
            has_degree = bool(_DEGREE_MODEL_RE.search(text))
            has_institution = bool(_INSTITUTION_MODEL_RE.search(text))
            if not has_degree and not has_institution and _YEAR_MODEL_RE.search(text):
                model.experiences.append(
                    Experience(
                        title=edu.degree or edu.school,
                        company=edu.school if edu.degree else "",
                        location=getattr(edu, "location", ""),
                        start_date=edu.start_date,
                        end_date=edu.end_date,
                    )
                )
            else:
                remaining_edu.append(edu)
        model.education = remaining_edu

    # ── 5. Language sanity (final pass) ──
    if hasattr(model, "languages") and model.languages:
        _tech_re = re.compile(
            r"\b(?:python|javascript|typescript|react|angular|docker|kubernetes"
            r"|sql|html|css|git|aws|azure|node\.?js)\b",
            re.I,
        )
        model.languages = [
            lang
            for lang in model.languages
            if lang
            and len(lang.strip()) > 1
            and not _tech_re.search(lang)
            and not re.match(r"https?://", lang, re.I)
            and not re.match(r"^[\d\W]+$", lang.strip())
            and "@" not in lang
        ]

    # ── 6. Misc cleanup: long misc → summary ──
    if hasattr(model, "misc") and model.misc:
        kept: list[str] = []
        for item in model.misc:
            if len(item.split()) >= 15:
                model.summary = f"{model.summary} {item}".strip() if model.summary else item
            else:
                kept.append(item)
        model.misc = kept

    # ── Deduplicate after moves ──
    seen_edu: set[str] = set()
    deduped_edu: list = []
    for edu in model.education:
        key = f"{edu.school.lower().strip()}|{edu.degree.lower().strip()}"
        if key not in seen_edu:
            seen_edu.add(key)
            deduped_edu.append(edu)
    model.education = deduped_edu

    seen_exp: set[str] = set()
    deduped_exp: list = []
    for exp in model.experiences:
        key = f"{exp.title.lower().strip()}|{exp.company.lower().strip()}"
        if key not in seen_exp:
            seen_exp.add(key)
            deduped_exp.append(exp)
    model.experiences = deduped_exp


# ── Anomaly detection (CVModel) ───────────────────────────────────────

_SKILLS_MODEL_MAX = 100
_MISC_MODEL_MAX = 10
_CONTACT_MODEL_MAX_LEN = 300
_LANG_MODEL_MAX = 15


def _anomaly_detection_model(model: CVModel) -> None:
    """Detect and fix abnormal CV structure on the CVModel.

    Only *moves* or *trims* existing data — never creates new content.
    """

    # ── 1. Skills overflow ──
    if len(model.skills) > _SKILLS_MODEL_MAX:
        overflow = model.skills[_SKILLS_MODEL_MAX:]
        model.skills = model.skills[:_SKILLS_MODEL_MAX]
        if hasattr(model, "misc"):
            for item in overflow:
                if len(item.split()) <= 4:
                    model.misc.append(item)
    if model.skills_categorized:
        for cat in list(model.skills_categorized):
            items = model.skills_categorized[cat]
            if len(items) > _SKILLS_MODEL_MAX:
                model.skills_categorized[cat] = items[:_SKILLS_MODEL_MAX]

    # ── 2. No education but degree exists elsewhere ──
    if not model.education:
        kept_exp: list = []
        for exp in model.experiences:
            text = f"{exp.title} {exp.company} {exp.location} {exp.start_date} {exp.end_date}"
            has_degree = bool(_DEGREE_MODEL_RE.search(text))
            has_institution = bool(_INSTITUTION_MODEL_RE.search(text))
            has_year = bool(_YEAR_MODEL_RE.search(text))
            if has_degree and has_institution and has_year:
                model.education.append(
                    Education(
                        degree=exp.title,
                        school=exp.company,
                        location=getattr(exp, "location", ""),
                        start_date=exp.start_date,
                        end_date=exp.end_date,
                    )
                )
            else:
                kept_exp.append(exp)
        model.experiences = kept_exp

        if not model.education and hasattr(model, "misc") and model.misc:
            kept_misc: list[str] = []
            for item in model.misc:
                if _DEGREE_MODEL_RE.search(item) and _YEAR_MODEL_RE.search(item):
                    model.education.append(Education(degree=item[:120]))
                else:
                    kept_misc.append(item)
            model.misc = kept_misc

    # ── 3. Languages invalid ──
    if hasattr(model, "languages") and model.languages:
        _tech_re = re.compile(
            r"\b(?:python|javascript|typescript|react|angular|docker|kubernetes"
            r"|sql|html|css|git|aws|azure|node\.?js)\b",
            re.I,
        )
        seen: set[str] = set()
        clean: list[str] = []
        for lang in model.languages:
            t = lang.strip()
            if not t or len(t) <= 1:
                continue
            has_proficiency_detail = bool(
                re.search(
                    r"\b(?:[ABC][12]|native|fluent|advanced|intermediate|beginner|proficient)\b",
                    t,
                    re.I,
                )
            )
            if len(t.split()) > 6 and not has_proficiency_detail:
                continue
            if _tech_re.search(t):
                continue
            if "@" in t or re.match(r"https?://", t, re.I):
                continue
            if re.match(r"^[\d\W]+$", t):
                continue
            key = re.sub(r"[\s\-\(\)]", "", t).lower()
            if key in seen:
                continue
            seen.add(key)
            clean.append(lang)
        model.languages = clean[:_LANG_MODEL_MAX]

    # ── 4. Contact too long ──
    if model.email and len(model.email) > _CONTACT_MODEL_MAX_LEN:
        model.email = model.email[:_CONTACT_MODEL_MAX_LEN]
    if model.phone and len(model.phone) > _CONTACT_MODEL_MAX_LEN:
        model.phone = model.phone[:_CONTACT_MODEL_MAX_LEN]
    if hasattr(model, "linkedin") and model.linkedin and len(model.linkedin) > _CONTACT_MODEL_MAX_LEN:
        model.linkedin = model.linkedin[:_CONTACT_MODEL_MAX_LEN]
    if model.location and len(model.location) > _CONTACT_MODEL_MAX_LEN:
        model.location = model.location[:_CONTACT_MODEL_MAX_LEN]
    for field in ("email", "phone"):
        val = getattr(model, field, "")
        if val and len(val.split()) > 8:
            setattr(model, field, "")

    # ── 5. Misc too large ──
    if hasattr(model, "misc") and model.misc and len(model.misc) > _MISC_MODEL_MAX:
        promoted: list[str] = []
        kept: list[str] = []
        for item in model.misc:
            if len(item.split()) >= 10:
                promoted.append(item)
            else:
                kept.append(item)
        if promoted:
            extra = " ".join(promoted)
            model.summary = f"{model.summary} {extra}".strip() if model.summary else extra
        model.misc = kept[:_MISC_MODEL_MAX]


# ── Fallback rendering ────────────────────────────────────────────────

_FALLBACK_MIN_SECTIONS = 2


def _needs_fallback_render(model: CVModel) -> bool:
    """Return True when the model is too empty to render a useful CV."""
    section_count = sum(
        [
            bool(model.summary),
            bool(model.experiences),
            bool(model.education),
            bool(model.skills or model.skills_categorized),
            bool(model.projects),
            bool(getattr(model, "certifications", None)),
            bool(model.languages),
            bool(getattr(model, "interests", None)),
        ]
    )
    total_text = (
        len(model.summary)
        + sum(len(e.title) + len(" ".join(e.bullets)) for e in model.experiences)
        + sum(len(e.degree) + len(e.school) for e in model.education)
        + sum(len(s) for s in model.skills)
        + sum(len(p.name) for p in model.projects)
    )
    misc_items = getattr(model, "misc", None) or []
    # No text at all and no misc
    if total_text == 0 and not misc_items:
        return True
    # Only misc has content
    if section_count == 0 and misc_items:
        return True
    # Too few sections and very little text
    if section_count < _FALLBACK_MIN_SECTIONS and total_text < 50:
        return True
    return False


def _build_fallback_model(cv_data: dict, model: CVModel) -> CVModel:
    """Re-populate model from raw cv_data when structured parse produced nothing.

    Scans every value in cv_data and pushes text into summary + misc
    so the rendered CV is never empty.  Never invents new data.
    """
    _SKIP = {
        "full_name",
        "title",
        "email",
        "phone",
        "location",
        "linkedin",
        "language",
        "section_titles",
        "format_hints",
        "contact",
    }

    collected: list[str] = []
    for key, value in cv_data.items():
        if key.startswith("_") or key in _SKIP:
            continue
        if isinstance(value, str) and value.strip():
            collected.append(value.strip())
        elif isinstance(value, list):
            for item in value:
                if isinstance(item, str) and item.strip():
                    collected.append(item.strip())
                elif isinstance(item, dict):
                    for v in item.values():
                        if isinstance(v, str) and v.strip():
                            collected.append(v.strip())
                        elif isinstance(v, list):
                            for sub in v:
                                if isinstance(sub, str) and sub.strip():
                                    collected.append(sub.strip())

    if not collected:
        return model

    # Deduplicate
    seen: set[str] = set()
    unique: list[str] = []
    for text in collected:
        key = text.lower().strip()
        if key not in seen:
            seen.add(key)
            unique.append(text)

    # Preserve existing contact fields
    if not model.full_name:
        model.full_name = _normalize_text(cv_data.get("full_name", "")).strip()
    if not model.email:
        model.email = _normalize_text(cv_data.get("email", "")).strip()
    if not model.phone:
        model.phone = _normalize_text(cv_data.get("phone", "")).strip()

    # First long block → summary
    if not model.summary:
        for i, text in enumerate(unique):
            if len(text.split()) >= 5:
                model.summary = _normalize_text(text)
                unique.pop(i)
                break

    # Remaining → misc
    misc = getattr(model, "misc", None)
    if misc is None:
        misc = []
    existing = {m.lower().strip() for m in misc}
    for text in unique:
        if text.lower().strip() not in existing:
            misc.append(text)
            existing.add(text.lower().strip())
    model.misc = misc

    return model


# ── Final layout normalization ─────────────────────────────────────────────

_CANONICAL_ORDER_MODEL = [
    "summary",
    "experience",
    "projects",
    "skills",
    "certifications",
    "education",
    "languages",
    "interests",
    "misc",
]

_SECTION_FIELD_MODEL = {
    "summary": "summary",
    "experience": "experiences",
    "projects": "projects",
    "education": "education",
    "skills": "skills_categorized",
    "certifications": "certifications",
    "languages": "languages",
    "interests": "interests",
    "misc": "misc",
}

_CONTACT_KEYS = ["full_name", "title", "email", "phone", "location", "linkedin"]
_SECTION_DICT_KEYS = [
    "summary",
    "experiences",
    "projects",
    "education",
    "skills",
    "skills_categorized",
    "certifications",
    "languages",
    "interests",
    "misc",
]
_META_KEYS = ["language", "section_titles"]


def _normalize_layout_model(model: CVModel) -> None:
    """Rebuild *section_titles* in canonical order, dropping empty sections."""
    old = dict(model.section_titles)
    ordered: dict[str, str] = {}

    for sec in _CANONICAL_ORDER_MODEL:
        field = _SECTION_FIELD_MODEL.get(sec, sec)
        val = getattr(model, field, None)
        if sec == "skills" and not val:
            val = getattr(model, "skills", None)
        has = bool(val.strip()) if isinstance(val, str) else bool(val)
        if has and sec in old:
            ordered[sec] = old[sec]

    model.section_titles = ordered


def _model_sanity_score(model: CVModel) -> int:
    """Return 0-3 score based on presence of primary sections."""
    score = 0
    if model.experiences:
        score += 1
    if model.education:
        score += 1
    if model.summary and model.summary.strip():
        score += 1
    return score


def _model_integrity_check(model: CVModel) -> None:
    """Ensure at least one primary section (experience, education, summary) exists.

    If none present, promote the longest misc entry or any long text to summary.
    """
    has_experience = bool(model.experiences)
    has_education = bool(model.education)
    has_summary = bool(model.summary and model.summary.strip())

    if has_experience or has_education or has_summary:
        return

    misc = getattr(model, "misc", None) or []
    if misc:
        best_idx = max(range(len(misc)), key=lambda i: len(misc[i]))
        best = misc[best_idx].strip()
        if len(best.split()) >= 5:
            model.summary = best
            misc.pop(best_idx)
            model.misc = misc
            return

    skills = getattr(model, "skills", None) or []
    for i, sk in enumerate(skills):
        if len(sk.split()) >= 10:
            model.summary = sk.strip()
            skills.pop(i)
            model.skills = skills
            return

    langs = getattr(model, "languages", None) or []
    for i, lang in enumerate(langs):
        if len(lang.split()) >= 10:
            model.summary = lang.strip()
            langs.pop(i)
            model.languages = langs
            return

    if misc:
        model.summary = " ".join(m.strip() for m in misc if m.strip())
        model.misc = []


def _reorder_data_dict(data: dict) -> dict:
    """Return *data* with keys in canonical layout order."""
    result: dict = {}
    for key in _CONTACT_KEYS + _SECTION_DICT_KEYS + _META_KEYS:
        if key in data:
            result[key] = data[key]
    for key in data:
        if key not in result:
            result[key] = data[key]
    return result
