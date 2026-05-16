from __future__ import annotations

import io
import os
from collections.abc import Callable

from fastapi import HTTPException, UploadFile


CLAMAV_ENABLED = os.getenv("CLAMAV_ENABLED", "0").lower() in ("1", "true", "yes")
CLAMAV_HOST = os.getenv("CLAMAV_HOST", "localhost")
CLAMAV_PORT = int(os.getenv("CLAMAV_PORT", "3310") or "3310")

PDF_CONTENT_TYPES = {"application/pdf", "application/octet-stream", "", None}
DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
TEXT_SUFFIXES = {".txt", ".md"}


def scan_upload_for_viruses(contents: bytes) -> None:
    """Scan uploaded file bytes with ClamAV when enabled."""

    if not CLAMAV_ENABLED:
        return

    try:
        import clamd
    except Exception:
        raise HTTPException(status_code=500, detail="Virus scanning backend unavailable")

    try:
        client = clamd.ClamdNetworkSocket(host=CLAMAV_HOST, port=CLAMAV_PORT)
        result = client.instream(io.BytesIO(contents))
    except Exception:
        raise HTTPException(status_code=500, detail="Virus scan failed")

    try:
        _name, (status, signature) = next(iter(result.items()))
    except Exception:
        status, signature = None, None

    if status != "OK":
        detail = "File failed virus scan"
        if signature:
            detail = f"Malware detected: {signature}"
        raise HTTPException(status_code=400, detail=detail)


def extract_pdf_text(contents: bytes) -> str:
    """Extract plain text from PDF bytes with layout-aware fallbacks."""

    errors: list[str] = []
    text_parts: list[str] = []

    try:
        import pdfplumber

        with pdfplumber.open(io.BytesIO(contents)) as pdf:
            for page_index, page in enumerate(pdf.pages, start=1):
                extracted = ""
                try:
                    extracted = page.extract_text(layout=True) or ""
                except TypeError:
                    extracted = page.extract_text() or ""
                if extracted.strip():
                    text_parts.append(f"\n--- Page {page_index} ---\n{extracted.strip()}")
    except Exception:
        errors.append("pdfplumber")

    if not text_parts:
        try:
            import PyPDF2

            pdf_reader = PyPDF2.PdfReader(io.BytesIO(contents))
            for page_index, page in enumerate(pdf_reader.pages, start=1):
                extracted = page.extract_text()
                if extracted and extracted.strip():
                    text_parts.append(f"\n--- Page {page_index} ---\n{extracted.strip()}")
        except Exception:
            errors.append("PyPDF2")

    if errors and not text_parts:
        raise HTTPException(status_code=400, detail="Invalid PDF file")

    text = "\n".join(text_parts).strip()
    if not text:
        raise HTTPException(status_code=400, detail="PDF contains no extractable text")
    return text


def extract_docx_text(contents: bytes) -> str:
    """Extract plain text from DOCX bytes, including simple table content."""

    try:
        from docx import Document
    except Exception:
        raise HTTPException(status_code=500, detail="DOCX parser is not available")

    try:
        document = Document(io.BytesIO(contents))
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid DOCX file")

    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            line = " | ".join(cell for cell in cells if cell)
            if line:
                parts.append(line)

    text = "\n".join(parts).strip()
    if not text:
        raise HTTPException(status_code=400, detail="DOCX contains no extractable text")
    return text


def extract_plain_text(contents: bytes) -> str:
    text = contents.decode("utf-8", errors="ignore").strip()
    if not text:
        raise HTTPException(status_code=400, detail="Text file contains no readable text")
    return text


def validate_pdf_upload(
    contents: bytes,
    content_type: str | None,
    *,
    virus_scanner: Callable[[bytes], None] = scan_upload_for_viruses,
) -> None:
    """Apply upload security checks for PDF files."""

    if content_type not in PDF_CONTENT_TYPES:
        raise HTTPException(status_code=400, detail="Only PDF files are allowed")
    if len(contents) > 5_000_000:
        raise HTTPException(status_code=400, detail="File too large (max 5MB)")
    if not contents.startswith(b"%PDF-"):
        raise HTTPException(status_code=400, detail="Invalid PDF file")

    try:
        virus_scanner(contents)
    except HTTPException:
        raise
    except Exception:
        if os.getenv("CLAMAV_ENABLED", "0").lower() in ("1", "true", "yes"):
            raise HTTPException(status_code=500, detail="Virus scanner error")


def extract_upload_text(
    contents: bytes,
    content_type: str | None = "",
    filename: str | None = "",
    *,
    max_size: int = 5_000_000,
    virus_scanner: Callable[[bytes], None] = scan_upload_for_viruses,
    pdf_extractor: Callable[[bytes], str] = extract_pdf_text,
    docx_extractor: Callable[[bytes], str] = extract_docx_text,
    text_extractor: Callable[[bytes], str] = extract_plain_text,
) -> str:
    """Extract text from supported CV/JD uploads: PDF, DOCX, or TXT."""

    if len(contents) > max_size:
        raise HTTPException(status_code=400, detail=f"File too large (max {max_size // 1_000_000}MB)")

    ctype = (content_type or "").split(";")[0].strip().lower()
    suffix = os.path.splitext(filename or "")[1].lower()

    try:
        virus_scanner(contents)
    except HTTPException:
        raise
    except Exception:
        if os.getenv("CLAMAV_ENABLED", "0").lower() in ("1", "true", "yes"):
            raise HTTPException(status_code=500, detail="Virus scanner error")

    if ctype == "application/pdf" or suffix == ".pdf" or contents.startswith(b"%PDF-"):
        if not contents.startswith(b"%PDF-"):
            raise HTTPException(status_code=400, detail="Invalid PDF file")
        return pdf_extractor(contents)

    if ctype == DOCX_CONTENT_TYPE or suffix == ".docx":
        if not contents.startswith(b"PK"):
            raise HTTPException(status_code=400, detail="Invalid DOCX file")
        return docx_extractor(contents)

    if suffix in TEXT_SUFFIXES or (ctype == "text/plain" and suffix in ("", *TEXT_SUFFIXES)):
        return text_extractor(contents)

    raise HTTPException(
        status_code=400,
        detail="Unsupported file type (use PDF, DOCX, or TXT)",
    )


async def resolve_job_description_text(
    job_description: str = "",
    jd_file: UploadFile | None = None,
    *,
    upload_text_extractor: Callable[..., str] = extract_upload_text,
) -> str:
    """Resolve JD text from direct input or uploaded file."""

    direct = (job_description or "").strip()
    if direct:
        return direct

    if jd_file is None:
        raise HTTPException(status_code=400, detail="Job description is required")

    contents = await jd_file.read()
    return upload_text_extractor(
        contents,
        jd_file.content_type,
        jd_file.filename,
        max_size=1_000_000,
    )
